"""
PGIM Deal Analysis Dashboard
=============================
Run with:  streamlit run frontend/app.py

Structure
---------
This file is the entire Streamlit front-end.  It is split into these parts:

  1. Imports & path setup
  2. Page config + CSS styling
  3. Utility helpers
       - load_deals()          — scan configs/ for all deal JSON files
       - load_config()         — read a single deal config
       - _list_ollama_models() — query Ollama for installed models
       - _ollama_base_url()    — read Ollama URL from any available config
       - _fmt_model()          — build "model — description" display string
       - _parse_model_name()   — extract raw model name from display string
       - _apply_model_overrides() — patch a config dict with sidebar model choices
       - _run_script()         — run a backend Python script with model overrides
       - _read_excel_preview() — load a formatted output Excel into a DataFrame
       - _show_results()       — display Excel preview + map + download buttons
       - _markdown_to_docx()   — convert markdown text to a .docx bytes object
       - _extract_text()       — extract plain text from an uploaded file
  4. Preview table helpers
       - _PREVIEW_ROWS         — ordered list of every config field shown in the editor
       - _fields_to_df()       — dict → editable DataFrame
       - _df_to_fields()       — edited DataFrame → dict (with type coercion)
  5. Session state initialisation
  6. Sidebar  — navigation radio + LLM model selectors + active model badges
  7. Render functions  — one function per page section
       - render_existing_deals()          — ROUTE A2: view/edit/delete existing deal
       - render_new_deal_form()          — ROUTE B: new deal wizard (2 steps)
       - render_comparable_analysis()    — ROUTE C: comps upload + AI search
       - render_investment_rationale()   — ROUTE D: rationale generation + refine
  8. Router  — reads sidebar selection and calls the correct render function

Session state keys used across this file
-----------------------------------------
  nd_step            : int  — which step of the New Deal wizard is active (1 or 2)
  nd_preview_fields  : dict — LLM-derived field values waiting for user review
  active_page        : str  — active nav item: new_deal | existing | overview | comps | rationale
  deal_nav           : str  — which existing deal is selected in the Existing Deals sidebar
  ao_deal_select     : str  — deal selectbox in the Analysis Output workspace
  sb_analysis_model  : str  — analysis model (⚙️ Shared Settings); also drives online search
  comp_deal          : str  — last selected deal in the Comps section
  comp_deal_select   : str  — selectbox state in render_comparable_analysis
  comp_type_radio    : str  — which comp type tab is active
  ir_deal_select     : str  — selectbox state in render_investment_rationale
  ir_notes           : str  — analyst notes text area
  ir_refresh         : bool — whether to force re-extraction of cached PDFs
  ir_refinement      : str  — supervisor feedback for the Refine panel
"""

import io
import json
import os
import subprocess
import sys
import tempfile
import time
from pathlib import Path

import pandas as pd
import streamlit as st

# ── Path setup ────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "backend"))


# ── Cloud secrets bootstrap (Streamlit Cloud only) ─────────────────────────────
# On Streamlit Cloud there is no persistent configs/shared_settings.json, so build
# it from st.secrets / env vars. Gated to the cloud (/mount/src path) so it is a
# no-op locally / on a server that already has a real shared_settings.json.
def _bootstrap_cloud_secrets() -> None:
    import json as _j
    # secret name -> shared_settings.json key
    _map = {
        "MAPBOX_TOKEN":    "mapbox_token",
        "GOOGLE_MAPS_KEY": "google_maps_key",
        "KAKAO_API_KEY":   "kakao_api_key",
        "OPENAI_API_KEY":  "openai_api_key",
        "ONEMAP_EMAIL":    "onemap_email",
        "ONEMAP_PASSWORD": "onemap_password",
    }
    keys = tuple(_map) + ("GEOCODING_PROVIDER",)
    vals = {}
    try:
        for k in keys:
            if k in st.secrets:
                vals[k] = str(st.secrets[k])
    except Exception:
        pass
    for k in keys:
        vals.setdefault(k, os.environ.get(k, ""))
    if not any(vals.values()):
        return
    for k, v in vals.items():          # expose to backend subprocesses
        if v and not os.environ.get(k):
            os.environ[k] = v
    cfgdir = ROOT / "configs"
    cfgdir.mkdir(exist_ok=True)
    ssp = cfgdir / "shared_settings.json"
    existing = {}
    if ssp.exists():
        try:
            existing = _j.loads(ssp.read_text(encoding="utf-8"))
        except Exception:
            existing = {}
    existing["geocoding_provider"] = (
        vals.get("GEOCODING_PROVIDER") or existing.get("geocoding_provider") or "mapbox")
    for _sk, _dk in _map.items():
        if vals.get(_sk):
            existing[_dk] = vals[_sk]
    try:
        ssp.write_text(_j.dumps(existing, indent=2), encoding="utf-8")
    except Exception:
        pass


if "/mount/" in str(ROOT):   # Streamlit Cloud only
    _bootstrap_cloud_secrets()

# ── Corporate proxy TLS fix (trust OS cert store; no-op without truststore) ────
from tools import corp_ssl  # noqa: F401  — must import before any HTTPS call

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PGIM Deal Analysis",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
/* ── Sidebar base ─────────────────────────────────────────── */
[data-testid="stSidebar"] { background-color: #1A3A5C; }
[data-testid="stSidebar"] * { color: #ffffff !important; }
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.2) !important; }

/* ── Sidebar form inputs — dark text on white/light backgrounds ── */
/* Text inputs, password fields, number inputs */
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea {
    color: #1A3A5C !important;
}
/* Selectbox — target every child of the baseweb select control
   (singleValue, placeholder, input are all deeply nested with hashed classes) */
[data-testid="stSidebar"] [data-baseweb="select"] *,
[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] * {
    color: #1A3A5C !important;
}
/* Selectbox dropdown popover options (rendered in a portal outside the sidebar) */
[data-baseweb="popover"] * {
    color: #1A3A5C !important;
}

/* ── Sidebar buttons — navy bg + white text; invert on hover/click ── */
[data-testid="stSidebar"] [data-testid="stButton"] > button {
    background-color : #1A3A5C !important;
    color            : #ffffff !important;
    border           : 1px solid rgba(255,255,255,0.35) !important;
    transition       : background-color 0.15s, color 0.15s !important;
}
[data-testid="stSidebar"] [data-testid="stButton"] > button:hover,
[data-testid="stSidebar"] [data-testid="stButton"] > button:active,
[data-testid="stSidebar"] [data-testid="stButton"] > button:focus {
    background-color : #ffffff !important;
    color            : #1A3A5C !important;
    border           : 1px solid #1A3A5C !important;
}
/* Keep the PRIMARY (active-page) nav button legible: lighter-navy fill, white
   text, and don't invert to white on hover/focus. Covers both the `kind`
   attribute and the stBaseButton-primary testid used by Streamlit 1.58. */
[data-testid="stSidebar"] button[kind="primary"],
[data-testid="stSidebar"] button[kind="primary"]:hover,
[data-testid="stSidebar"] button[kind="primary"]:active,
[data-testid="stSidebar"] button[kind="primary"]:focus,
[data-testid="stSidebar"] [data-testid="stBaseButton-primary"],
[data-testid="stSidebar"] [data-testid="stBaseButton-primary"]:hover,
[data-testid="stSidebar"] [data-testid="stBaseButton-primary"]:active,
[data-testid="stSidebar"] [data-testid="stBaseButton-primary"]:focus {
    background-color : #2E5A85 !important;
    color            : #ffffff !important;
    border           : 1px solid #ffffff !important;
}

/* ── Sidebar expander headers (Shared Settings / nav groups) ──
   Keep the dark sidebar background; never flash white on click / focus /
   open. Only a subtle translucent tint on hover. */
[data-testid="stSidebar"] [data-testid="stExpander"] {
    background-color: transparent !important;
    border: 1px solid rgba(255,255,255,0.18) !important;
}
[data-testid="stSidebar"] [data-testid="stExpander"] details,
[data-testid="stSidebar"] [data-testid="stExpander"] details[open],
[data-testid="stSidebar"] [data-testid="stExpander"] summary,
[data-testid="stSidebar"] [data-testid="stExpander"] summary:focus,
[data-testid="stSidebar"] [data-testid="stExpander"] summary:active,
[data-testid="stSidebar"] [data-testid="stExpander"] details[open] > summary,
[data-testid="stSidebar"] [data-testid="stExpanderHeader"],
[data-testid="stSidebar"] .streamlit-expanderHeader {
    background-color: transparent !important;
    color: #ffffff !important;
    box-shadow: none !important;
}
[data-testid="stSidebar"] [data-testid="stExpander"] summary:hover,
[data-testid="stSidebar"] [data-testid="stExpanderHeader"]:hover,
[data-testid="stSidebar"] .streamlit-expanderHeader:hover {
    background-color: rgba(255,255,255,0.10) !important;
    color: #ffffff !important;
}

/* ── Sidebar radio as nav menu ───────────────────────────── */
[data-testid="stSidebar"] .stRadio > label { display: none; }
[data-testid="stSidebar"] .stRadio > div   { gap: 4px; }
[data-testid="stSidebar"] .stRadio label {
    background    : rgba(255,255,255,0.08);
    border-radius : 8px;
    padding       : 9px 14px !important;
    margin        : 0 !important;
    cursor        : pointer;
    font-size     : 14px !important;
    font-weight   : 500 !important;
    width         : 100%;
    transition    : background 0.15s;
}
[data-testid="stSidebar"] .stRadio label:hover {
    background: rgba(255,255,255,0.18);
}
/* deal sub-list items smaller */
[data-testid="stSidebar"] .stRadio.deal-list label {
    font-size  : 13px !important;
    font-weight: 400 !important;
    padding    : 6px 14px 6px 22px !important;
    background : transparent;
}
[data-testid="stSidebar"] .stRadio.deal-list label:hover {
    background: rgba(255,255,255,0.10);
}

/* ── Main headings ────────────────────────────────────────── */
h1, h2, h3 { color: #1A3A5C; }
.stTabs [data-baseweb="tab"] { font-size: 14px; font-weight: 600; }
.stTabs [data-baseweb="tab-highlight"] { background-color: #1A3A5C; }

/* ── Deal summary card ────────────────────────────────────── */
.deal-card {
    background    : #f7f9fc;
    border        : 1px solid #dce6f0;
    border-radius : 10px;
    padding       : 18px 22px;
    margin-bottom : 16px;
}
.section-label {
    font-size   : 11px;
    font-weight : 700;
    letter-spacing: 1px;
    color       : #1A3A5C;
    text-transform: uppercase;
    margin-bottom: 8px;
}
</style>
""", unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def load_deals() -> dict:
    """Return {deal_name: config_path} for all saved configs."""
    deals = {}
    for f in sorted((ROOT / "configs").glob("deal_config*.json")):
        try:
            cfg  = json.loads(f.read_text(encoding="utf-8"))
            name = cfg["subject_property"]["deal_name"]
            deals[name] = str(f)
        except Exception:
            pass
    return deals


def load_config(path: str) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


@st.cache_data(ttl=15)
def _list_ollama_models(base_url: str = "http://localhost:11434") -> list[str]:
    """Query Ollama for installed models. Returns [] if Ollama is not running."""
    try:
        import urllib.request
        with urllib.request.urlopen(f"{base_url}/api/tags", timeout=2) as r:
            data = json.loads(r.read())
        names = [m["name"] for m in data.get("models", [])]
        return sorted(names) if names else []
    except Exception:
        return []


_OPENAI_ANALYSIS_MODELS  = ["gpt-4o-mini", "gpt-4o"]   # fast → best quality
_OPENAI_SEARCH_MODELS    = ["gpt-4o-mini-search-preview", "gpt-4o-search-preview"]
_OPENAI_VISION_MODELS    = ["gpt-4o-mini", "gpt-4o"]   # both support vision

# Analysis-model option that turns the LLM off entirely: comps are classified by
# keyword rules and columns mapped by exact synonyms only — no Ollama, no cloud.
# Use it to run the pipeline (e.g. to test OneMap geocoding) without any LLM.
_NO_LLM_OPTION = "🚫 Rule-based (no LLM)"

# Ollama model name substrings that indicate vision capability.
# Used to filter the installed model list for the vision model selector.
_VISION_MODEL_KEYWORDS   = {"llava", "minicpm", "vision", "bakllava", "moondream"}

# Currency options for the New Deal form.
# Key   = display label shown in the dropdown
# Value = (currency_code, currency_symbol) written into the deal config
_CURRENCY_OPTIONS: dict[str, tuple] = {
    "Auto-detect from country":  ("", ""),
    "SGD — Singapore Dollar":    ("SGD", "S$"),
    "USD — US Dollar":           ("USD", "$"),
    "AUD — Australian Dollar":   ("AUD", "A$"),
    "JPY — Japanese Yen":        ("JPY", "¥"),
    "KRW — South Korean Won":    ("KRW", "₩"),
    "HKD — Hong Kong Dollar":    ("HKD", "HK$"),
    "GBP — British Pound":       ("GBP", "£"),
    "EUR — Euro":                ("EUR", "€"),
    "CNY — Chinese Yuan":        ("CNY", "¥"),
    "INR — Indian Rupee":        ("INR", "₹"),
}

# Preferred display order for local models — fastest (smallest) to slowest (largest).
# Any installed model NOT in this list is appended alphabetically at the end.
_MODEL_SPEED_ORDER = [
    "qwen2.5:3b",       # 1.9 GB  — fastest
    "gemma3:4b",        # 3.3 GB
    "deepseek-r1:7b",   # 4.7 GB
    "llama3.1:8b",      # 4.9 GB
    "qwen3:8b",         # 5.2 GB
    "qwen3.5:9b",       # ~6 GB   — slowest local
]

# Descriptions shown next to each model name in the sidebar selectbox.
# Ordered to match _MODEL_SPEED_ORDER so the comment column is easy to scan.
_MODEL_DESCRIPTIONS: dict[str, str] = {
    # ── Local models (fast → slow) ────────────────────────────────────────────
    "qwen2.5:3b":       "⚡ Fastest — lightweight tasks",
    "gemma3:4b":        "⚡ Fast — decent quality",
    "deepseek-r1:7b":   "🧠 Investment reasoning & analysis",
    "llama3.1:8b":      "✍️  General purpose, good writing",
    "qwen3:8b":         "🏦 Finance + Asian markets",
    "qwen3.5:9b":       "🏦 Finance + Asian markets, stronger reasoning",
    # ── Cloud models (OpenAI) ─────────────────────────────────────────────────
    "gpt-4o-mini":      "☁️  Fast cloud — investment rationale writing",
    "gpt-4o":           "☁️  Best cloud quality",
}


def _sort_models_by_speed(models: list[str]) -> list[str]:
    """
    Sort a list of installed Ollama model names by preferred speed order.
    Models in _MODEL_SPEED_ORDER come first (in that order).
    Any unlisted model is appended alphabetically at the end.
    """
    known   = [m for m in _MODEL_SPEED_ORDER if m in models]
    unknown = sorted(m for m in models if m not in _MODEL_SPEED_ORDER)
    return known + unknown


def _fmt_model(name: str) -> str:
    """Return 'model — description' if a description exists, else just 'model'."""
    desc = _MODEL_DESCRIPTIONS.get(name, "")
    return f"{name} — {desc}" if desc else name


def _parse_model_name(display_str: str) -> str:
    """Extract just the model name from a display string like 'model — description'."""
    return display_str.split(" — ")[0].strip() if " — " in display_str else display_str.strip()


def _apply_model_overrides(cfg: dict) -> dict:
    """
    Patch a config dict with the LLM model selections from the sidebar.
    Returns a (possibly modified) deep copy — original is untouched.
    Session state may hold display strings like 'model — description';
    we always parse out just the model name before applying.
    """
    analysis_model = _parse_model_name(st.session_state.get("sb_analysis_model", ""))
    _raw_vision    = st.session_state.get("sb_vision_model", "")
    vision_model   = ("" if _raw_vision.startswith("──")
                      else _parse_model_name(_raw_vision))

    if not analysis_model and not vision_model:
        return cfg  # nothing to patch

    import copy
    cfg = copy.deepcopy(cfg)

    if analysis_model:
        if analysis_model == _NO_LLM_OPTION:
            # Rule-based only — backend skips the LLM entirely.
            cfg.setdefault("llm", {})["provider"] = "none"
        elif analysis_model in _OPENAI_ANALYSIS_MODELS:
            cfg.setdefault("llm", {})["provider"]       = "openai"
            cfg.setdefault("llm", {})["openai_model"]   = analysis_model
            cfg.setdefault("llm", {})["openai_api_key"] = os.environ.get("OPENAI_API_KEY", "")
        else:
            cfg.setdefault("llm", {})["provider"] = "ollama"
            cfg.setdefault("llm", {}).setdefault("ollama", {})["model"] = analysis_model

        # Online search is folded into the analysis-model choice. Web search only
        # runs on OpenAI search-preview models (Ollama can't browse), so map the
        # analysis model to the matching search + extract models: gpt-4o → the
        # full search model, everything else (mini / Ollama / rule-based) → the
        # cheaper mini search model.
        if analysis_model == "gpt-4o":
            cfg.setdefault("openai", {})["search_model"]  = "gpt-4o-search-preview"
            cfg.setdefault("openai", {})["extract_model"] = "gpt-4o"
        else:
            cfg.setdefault("openai", {})["search_model"]  = "gpt-4o-mini-search-preview"
            cfg.setdefault("openai", {})["extract_model"] = "gpt-4o-mini"

    if vision_model:
        if vision_model in _OPENAI_VISION_MODELS:
            cfg.setdefault("llm", {})["openai_vision_model"] = vision_model
        else:
            # Ollama vision model — stored separately so text tasks keep using
            # the fast analysis model while image tasks use the vision model.
            cfg.setdefault("llm", {}).setdefault("ollama", {})["vision_model"] = vision_model

    return cfg


def _run_script(script: str, config_path: str, flags: list = None,
                expand_log: bool = False, log_state_key: str = None,
                stream_live: bool = True) -> bool:
    """Run a backend script, injecting the sidebar LLM model overrides into the config.

    ``stream_live`` — when True (default), the last 40 log lines stream into a live
    box during the run. Set False (e.g. on the Overview) to show only the one-line
    'Running … [model] …' spinner; the full log stays in the collapsed run-log
    expander for manual inspection."""
    cfg      = load_config(config_path)
    patched  = _apply_model_overrides(cfg)

    use_path    = config_path
    tmp_path    = None
    if patched is not cfg:          # overrides were applied → write a temp config
        tmp = tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", dir=ROOT / "configs", delete=False,
            encoding="utf-8",
        )
        json.dump(patched, tmp)
        tmp.close()
        tmp_path = tmp.name
        use_path = tmp_path

    cmd = [sys.executable, str(ROOT / "backend" / script),
           "--config", use_path] + (flags or [])
    active_model = _parse_model_name(st.session_state.get("sb_analysis_model", ""))
    model_label  = f" [{active_model}]" if active_model else ""
    # Force the child pipeline to emit UTF-8 on its (piped) stdout/stderr. On Windows
    # a piped stdout defaults to the locale encoding (cp1252/GBK), so printing Korean
    # names or "㎡" would raise UnicodeEncodeError inside the child before we capture it.
    # PYTHONUNBUFFERED=1 makes the child flush each print() immediately, so the run log
    # below streams live instead of appearing all at once when the script finishes.
    _child_env = {**os.environ, "PYTHONUTF8": "1", "PYTHONIOENCODING": "utf-8",
                  "PYTHONUNBUFFERED": "1"}

    # Stream stdout+stderr line-by-line into a live-updating box so long-running
    # scripts (e.g. Online Search) show progress instead of a silent spinner.
    _lines, _live = [], st.empty()
    _spin = (f"Running `{script}`{model_label} … (live log below)" if stream_live
             else f"Running `{script}`{model_label} …")
    with st.spinner(_spin):
        proc = subprocess.Popen(
            cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True,
            encoding="utf-8", errors="replace", cwd=str(ROOT), env=_child_env,
            bufsize=1,
        )
        for _ln in proc.stdout:
            _lines.append(_ln.rstrip("\n"))
            if stream_live:
                _live.code("\n".join(_lines[-40:]), language="text")   # last 40 lines, snappy
        proc.wait()
    _live.empty()   # replace the live tail with the full expander below
    out, returncode = "\n".join(_lines), proc.returncode

    # Human flag: comps that geocoded to the country centroid are almost certainly
    # mislocated — surface a prominent warning above the run log.
    _n_suspect = out.count("ON COUNTRY CENTROID")
    if _n_suspect:
        st.warning(f"⚠ {_n_suspect} comparable(s) geocoded to the **country centroid** — "
                   "likely invalid coordinates (they'll pile up at one point on the map). "
                   "Check the run log and verify/correct those addresses before using.")

    if tmp_path:
        try: os.unlink(tmp_path)
        except: pass

    _key = log_state_key or f"_runlog_{script}"
    if out:
        st.session_state[_key] = out
        with st.expander("📋 Run log", expanded=expand_log):
            st.code(out, language="text")
    if returncode != 0:
        st.session_state[f"{_key}_err"] = out
        with st.expander("⚠️ Error details", expanded=True):
            st.code(out[-4000:] or "(no output)", language="text")
    return returncode == 0


def _read_excel_preview(excel_path: str) -> "pd.DataFrame | None":
    """
    Read the formatted output Excel and return a clean DataFrame for preview.
    Handles the two-table layout (subject row + comps section).
    """
    try:
        import openpyxl
        from openpyxl.utils import get_column_letter
        wb   = openpyxl.load_workbook(excel_path, data_only=True)
        ws   = wb.active
        # Find hidden column indices (stake % helper column etc.)
        hidden = {
            i for i in range(1, ws.max_column + 1)
            if ws.column_dimensions.get(get_column_letter(i),
               type("_", (), {"hidden": False})()).hidden
        }
        rows = [
            tuple(c.value for i, c in enumerate(row, 1) if i not in hidden)
            for row in ws.iter_rows()
        ]

        # Find first header row (has ≥3 cells + contains 'property' + 'marker')
        header_row, header_idx = None, None
        for i, row in enumerate(rows):
            strs = [str(c or "").lower() for c in row]
            if (sum(1 for s in strs if s) >= 3
                    and any("property" in s for s in strs)
                    and any("marker" in s for s in strs)):
                header_row, header_idx = row, i
                break

        if header_row is None:
            return None

        # Clean column headers — strip newlines
        headers = [str(h or "").replace("\n", " ").strip() for h in header_row]
        prop_idx = next((i for i, h in enumerate(headers) if "property" in h.lower()), None)

        # Deduplicate headers — keep first occurrence, drop later duplicates.
        # Guards against Excel files generated with a buggy schema that wrote
        # the same column header twice.
        _seen_h: set = set()
        _keep_cols = [i for i, h in enumerate(headers)
                      if h and h not in _seen_h and not _seen_h.add(h)]
        if len(_keep_cols) < len(headers):
            headers = [headers[i] for i in _keep_cols]

        data_rows = []
        for row in rows[header_idx + 1:]:
            # Skip fully empty rows
            if all(c in (None, "") for c in row):
                continue
            strs = [str(c or "").lower() for c in row]
            # Skip duplicate header rows (second table header)
            if (any("property" in s for s in strs)
                    and any("marker" in s for s in strs)):
                continue
            # Skip notes/footer rows
            if str(row[0] or "").strip().lower().startswith("note"):
                continue
            # Skip Average summary row (computed, not editable data)
            if any(str(v or "").strip().lower() == "average" for v in row):
                continue
            data_rows.append([row[i] for i in _keep_cols if i < len(row)])

        if not data_rows:
            return None

        df = pd.DataFrame(data_rows, columns=headers)
        # Drop only columns with no header (truly unnamed/invisible helper columns).
        # Do NOT drop columns just because all values are None — formula columns
        # (Price psf, Adj Cap Rate) legitimately have None values when read by
        # openpyxl, but should still appear in the preview with "—" placeholders.
        df = df.loc[:, [bool(h.strip()) for h in df.columns]]
        df = df.fillna("—").astype(str).replace("None", "—")
        return df

    except Exception:
        return None


def _save_edited_preview(excel_path: str, edited_df: "pd.DataFrame") -> bool:
    """
    Write the user-edited preview DataFrame back into the formatted Excel file.

    Strategy: locate the comp-data header row in the active sheet, clear every
    visible data cell below it, then write the new values from edited_df.
    Cell formatting (colours, fonts, borders) is preserved because we only
    change .value — openpyxl does not touch styles unless you explicitly set them.

    Returns True on success, False on any error.
    """
    try:
        import openpyxl
        from openpyxl.utils import get_column_letter

        wb = openpyxl.load_workbook(excel_path)
        ws = wb.active

        # Identify hidden columns (same logic as _read_excel_preview)
        hidden = {
            i for i in range(1, ws.max_column + 1)
            if ws.column_dimensions.get(get_column_letter(i),
               type("_", (), {"hidden": False})()).hidden
        }
        visible = [i for i in range(1, ws.max_column + 1) if i not in hidden]

        # Find the second header row — the comps table header (not the subject one).
        # The Excel has two header rows both containing "property" and "marker":
        # the first is for the subject section, the second is for the comps section.
        # We must only rewrite the comp data area (below the second header) so the
        # subject row and the fixed section-header rows between the two tables are
        # left completely untouched.
        headers_found = 0
        comp_header_row_idx = None
        for row in ws.iter_rows():
            vals = [str(c.value or "").lower() for c in row]
            if (sum(1 for v in vals if v) >= 3
                    and any("property" in v for v in vals)
                    and any("marker" in v for v in vals)):
                headers_found += 1
                if headers_found == 2:
                    comp_header_row_idx = row[0].row
                    break

        # Fallback: if only one header exists (single-table layouts), scan for
        # the first header instead so the function still works.
        if comp_header_row_idx is None:
            for row in ws.iter_rows():
                vals = [str(c.value or "").lower() for c in row]
                if (sum(1 for v in vals if v) >= 3
                        and any("property" in v for v in vals)
                        and any("marker" in v for v in vals)):
                    comp_header_row_idx = row[0].row
                    break

        if comp_header_row_idx is None:
            return False

        # Filter to comp rows only — strip out the subject row (★) and any
        # Average summary row so they are never written into the comp data area.
        marker_col = next((h for h in edited_df.columns if "marker" in h.lower()), None)
        prop_col   = next((h for h in edited_df.columns if "property" in h.lower()), None)
        comp_df = edited_df
        if marker_col:
            comp_df = comp_df[comp_df[marker_col].astype(str).str.strip() != "★"]
        if prop_col:
            comp_df = comp_df[~comp_df[prop_col].astype(str).str.strip()
                               .str.lower().str.startswith("average")]

        # Scan for any existing Average row in the comp area so the clear range
        # always covers it, even when ws.max_row under-counts styled-but-empty cells.
        old_avg_row = None
        scan_end = max(ws.max_row, comp_header_row_idx + 300) + 1
        for r in range(comp_header_row_idx + 1, scan_end):
            for col_idx in visible:
                if str(ws.cell(row=r, column=col_idx).value or "").strip().lower() == "average":
                    old_avg_row = r
                    break
            if old_avg_row:
                break

        # Clear all existing data rows in the comp area, ensuring the old Average
        # row is always included in the cleared range.
        clear_end = max(ws.max_row, old_avg_row or 0) + 1

        # Unmerge any merged ranges overlapping the comp area before writing —
        # a MergedCell's .value is read-only, so clearing/rewriting one raises
        # "'MergedCell' object attribute 'value' is read-only". We rewrite the
        # whole area anyway, so dropping these merges is safe.
        write_end = comp_header_row_idx + 1 + len(comp_df) + 2  # comp rows + Average + buffer
        unmerge_end = max(clear_end, write_end)
        for rng in list(ws.merged_cells.ranges):
            if rng.max_row >= comp_header_row_idx + 1 and rng.min_row <= unmerge_end:
                ws.unmerge_cells(str(rng))

        for r in range(comp_header_row_idx + 1, clear_end):
            for col_idx in visible:
                ws.cell(row=r, column=col_idx).value = None

        # Write edited rows — convert "—" / "None" placeholders back to None.
        # Coerce numpy/pandas scalar types to plain Python so openpyxl never
        # receives a numpy.int64 / numpy.float64 — those trip openpyxl's
        # internal style handling which expects str or native Python numerics.
        def _coerce(v):
            sv = str(v).strip()
            if sv in ("—", "None", "nan", ""):
                return None
            if hasattr(v, "item"):
                return v.item()
            return v

        n_data = 0
        for i, (_, row_data) in enumerate(comp_df.iterrows()):
            excel_row = comp_header_row_idx + 1 + i
            for j, val in enumerate(row_data):
                if j >= len(visible):
                    break
                ws.cell(row=excel_row, column=visible[j]).value = _coerce(val)
            n_data += 1

        # Re-write the Average row immediately after the last comp row.
        # Uses AVERAGEIF(col, ">0") to skip blank / zero cells.
        avg_excel_row = comp_header_row_idx + 1 + n_data
        header_vals = [str(ws.cell(row=comp_header_row_idx, column=col).value or "").lower().replace("\n", " ").strip()
                       for col in visible]
        _avg_targets = {
            "price": '#,##0',
            "psf":   '#,##0',
            "cap":   '0.00%',
            "yield": '0.00%',
        }
        first_r = comp_header_row_idx + 1
        last_r  = comp_header_row_idx + n_data
        for j, col_idx in enumerate(visible):
            hdr = header_vals[j] if j < len(header_vals) else ""
            cell = ws.cell(row=avg_excel_row, column=col_idx)
            if "property" in hdr:
                cell.value = "Average"
            elif any(kw in hdr for kw in _avg_targets) and first_r <= last_r:
                col_ltr = get_column_letter(col_idx)
                cell.value = (f'=IFERROR(AVERAGEIF({col_ltr}{first_r}:'
                              f'{col_ltr}{last_r},">0"),"—")')
            else:
                cell.value = None

        wb.save(excel_path)
        return True
    except Exception as e:
        import traceback
        st.error(f"Save failed: {e}\n\n```\n{traceback.format_exc()}\n```")
        return False


def _sync_records_json(records_path, edited_df: "pd.DataFrame", marker_map: dict = None):
    """
    Keep _records.json in sync with the edited preview table after a Save.

    Rules applied (all matched by Map Marker):
      • Row deleted from table  → record removed from JSON
      • Cell edited in table    → corresponding raw field updated in JSON
      • New row (no marker)     → appended to JSON with _source='manual'
                                  (will be geocoded on next Re-run)

    `edited_df` must carry the ORIGINAL map markers (matching _records.json) so
    deletions/edits resolve correctly.  `marker_map` (old → new) renumbers the
    surviving records afterwards to match the table's sequential renumbering.

    Only fields that have a direct column→key mapping are touched;
    calculated fields (Price PSF, Adj Cap Rate) are left alone.
    """
    import json as _json
    from pathlib import Path as _Path

    records_path = _Path(records_path)
    if not records_path.exists():
        return

    try:
        records = _json.loads(records_path.read_text(encoding="utf-8"))
    except Exception:
        return

    # Output column (lower-stripped) → raw record key
    _COL_MAP = {
        "sale date":               "sale_date",
        "land zoning":             "land_zoning",
        "remaining leasehold (y)": "remaining_yrs",
        "remaining leasehold":     "remaining_yrs",
        "gfa (sf)":                "gfa_sf",
        "price (sgd m)":           "price_sgd_m",
        "ftm noi cap rate":        "npi_yield",
        "sale type":               "sale_type",
        "asset type":              "asset_type",
    }

    # Normalised column name → actual DataFrame column name
    df_norm = {c.lower().strip(): c for c in edited_df.columns}

    marker_col = next((c for c in edited_df.columns if "marker"   in c.lower()), None)
    prop_col   = next((c for c in edited_df.columns if "property" in c.lower()), None)
    addr_col   = next((c for c in edited_df.columns if "address"  in c.lower()), None)

    # Markers still present in the edited table
    remaining = set()
    if marker_col:
        remaining = {
            str(v).strip()
            for v in edited_df[marker_col].dropna()
            if str(v).strip() not in ("", "—", "None", "★", "*")
        }

    # Detect whether records have real map_marker values
    _has_markers = any(r.get("map_marker") for r in records)

    if _has_markers:
        # 1a. Remove records whose marker is no longer in the table
        records = [r for r in records if str(r.get("map_marker", "")) in remaining]
        idx_by_marker = {str(r.get("map_marker", "")): i for i, r in enumerate(records)}
    else:
        # 1b. Fallback: records have no map_marker yet — keep all and match by name
        idx_by_marker = {}
        if prop_col:
            for i, r in enumerate(records):
                name = str(r.get("property_name") or r.get("site_name") or "").strip()
                if name:
                    idx_by_marker[name.lower()] = i

    def _parse_val(v: str, key: str):
        """Convert a display string to the correct Python type for *key*."""
        v = v.replace(",", "").replace("%", "").strip()
        if not v or v in ("—", "None", "nan"):
            return None
        if key == "gfa_sf":
            try:    return int(float(v))
            except: return None
        if key == "remaining_yrs":
            try:    return round(float(v), 1)
            except: return None
        if key in ("price_sgd_m", "npi_yield"):
            try:    return float(v)
            except: return None
        return v  # string field

    # 2. Apply edits to existing records / append new rows
    for _, row in edited_df.iterrows():
        marker = str(row.get(marker_col, "") if marker_col else "").strip()
        if _has_markers:
            is_existing = marker and marker not in ("—", "None", "★", "*") and marker in idx_by_marker
            rec_idx = idx_by_marker.get(marker)
        else:
            # Name-based fallback
            raw_prop_val = str(row.get(prop_col, "") if prop_col else "").strip()
            prop_lines   = [ln.strip() for ln in raw_prop_val.split("\n") if ln.strip()]
            prop_name_lc = (prop_lines[0] if prop_lines else "").lower()
            rec_idx      = idx_by_marker.get(prop_name_lc)
            is_existing  = rec_idx is not None

        if is_existing and rec_idx is not None:
            rec = records[rec_idx]
            for norm, raw_key in _COL_MAP.items():
                if norm in df_norm:
                    parsed = _parse_val(str(row.get(df_norm[norm], "")), raw_key)
                    if parsed is not None:
                        rec[raw_key] = parsed

            # Persist Property-name edits (kept as a raw string — not run through
            # the numeric cleaner) so corrections like 'GIEO'→'GEO' survive a
            # re-run, which reads these records rather than re-parsing the PDF.
            if prop_col:
                new_name = str(row.get(prop_col, "")).replace("\n", " ").strip()
                if new_name and new_name not in ("—", "None", "nan"):
                    name_key = next((k for k in ("property_name", "site_name",
                                                 "building_name") if k in rec),
                                    "property_name")
                    if new_name != str(rec.get(name_key, "")):
                        rec[name_key] = new_name
                        # invalidate cached coords so the re-run re-geocodes it
                        rec.pop("lon", None); rec.pop("lat", None)

            # Persist Address edits — when the analyst types an address for a
            # property that couldn't be geocoded by name, save it so the re-run
            # geocodes by address (address takes priority over name).
            if addr_col:
                new_addr = str(row.get(addr_col, "")).replace("\n", " ").strip()
                if new_addr in ("—", "None", "nan"):
                    new_addr = ""
                if new_addr != str(rec.get("address", "") or ""):
                    rec["address"] = new_addr
                    rec.pop("lon", None); rec.pop("lat", None)  # re-geocode w/ new address

        else:
            # New row — build a minimal record for _source='manual'
            if not prop_col:
                continue
            raw_prop = str(row.get(prop_col, "")).strip()
            lines    = [ln.strip() for ln in raw_prop.split("\n") if ln.strip()]
            name     = lines[0] if lines else ""
            addr     = lines[1] if len(lines) > 1 else ""
            # Prefer an explicit Address-column value if the analyst filled one.
            if addr_col:
                _a = str(row.get(addr_col, "")).replace("\n", " ").strip()
                if _a and _a not in ("—", "None", "nan"):
                    addr = _a
            if not name or name in ("—", "None"):
                continue

            new_rec = {
                "property_name":   name,
                "address":         addr,
                "raw_description": f"{name}\n{addr}".strip(),
                "_source":         "manual",
                "stake_pct":       1.0,
            }
            for norm, raw_key in _COL_MAP.items():
                if norm in df_norm:
                    parsed = _parse_val(str(row.get(df_norm[norm], "")), raw_key)
                    if parsed is not None:
                        new_rec[raw_key] = parsed

            if not new_rec.get("price_sgd_m"):
                continue   # price required
            records.append(new_rec)

    # Renumber surviving records' map_marker to match the table's sequential
    # renumbering done at save time (old → new).
    if marker_map and _has_markers:
        for rec in records:
            om = str(rec.get("map_marker", ""))
            if om in marker_map:
                rec["map_marker"] = marker_map[om]

    try:
        records_path.write_text(
            _json.dumps(records, indent=2, default=str), encoding="utf-8"
        )
    except Exception:
        pass


def _show_interactive_map(geo_path, excel_path, context: str):
    """
    Render an interactive pydeck map from the _geo.json sidecar.

    Each pin shows a hover tooltip with property name + address.
    A 'Manage Pins' panel below lets the analyst:
      - Delete pin only  → sets hidden=true in geo JSON, regenerates PNG
      - Delete pin + row → sets hidden=true in geo JSON AND deletes Excel row, regenerates PNG
      - Restore          → sets hidden=false in geo JSON, regenerates PNG

    The 'hidden' flag in the geo JSON is the single source of truth.
    Session state is NOT used for pin visibility.
    """
    import json as _json
    import pydeck as pdk

    geo_path = Path(geo_path)
    if not geo_path.exists():
        return False

    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
    except Exception:
        return False

    subj     = geo.get("subject", {})
    s_lon    = subj.get("lon")
    s_lat    = subj.get("lat")
    if s_lon is None:
        return False

    all_comps = geo.get("comps", [])
    mb        = geo.get("mapbox", {})
    token     = _get_mapbox_token()   # always from shared_settings.json

    # ── Build layer data — read 'hidden' flag directly from geo JSON ──────────
    import math as _math
    visible = [
        {
            "lon":      c["lon"],
            "lat":      c["lat"],
            "property": c.get("property", ""),
            "address":  c.get("address", ""),
            "marker":   str(c["map_marker"]),
            "color":    c.get("color"),
        }
        for c in all_comps
        if c.get("lon") is not None and not c.get("hidden", False)
    ]

    _show_subject = st.checkbox(
        "📍 Plot subject property (★) on map", value=True,
        key=f"map_show_subject_{context}_{geo_path.stem}",
        help="Untick to hide the subject star — affects both this map and the "
             "downloadable PNG.",
    )
    # Keep the downloadable static PNG in sync with the checkbox. Regenerate only
    # when the value changes (the run's PNG already includes the subject = True).
    _subj_png_key = f"_subj_png_state_{context}_{geo_path.stem}"
    if _subj_png_key not in st.session_state:
        st.session_state[_subj_png_key] = True
    if st.session_state[_subj_png_key] != _show_subject:
        _regen_map_from_geo(geo_path, plot_subject=_show_subject)
        st.session_state[_subj_png_key] = _show_subject
    subject_pt = [{"lon": s_lon, "lat": s_lat,
                   "property": "Subject property", "address": "", "marker": "★"}]

    # Centre view on the plotted points and fit zoom to their spread.
    # When the subject is hidden, centre on the comps only (fall back to the
    # subject just to keep a valid view if there are no comps).
    import math as _math
    plotted_pts = visible + (subject_pt if _show_subject else [])
    all_pts = plotted_pts or subject_pt or visible
    center_lon = sum(p["lon"] for p in all_pts) / len(all_pts)
    center_lat = sum(p["lat"] for p in all_pts) / len(all_pts)

    def _fit_zoom(pts, w=900, h=520, pad=0.35):
        if len(pts) <= 1:
            return 14
        lons = [p["lon"] for p in pts]
        lats = [p["lat"] for p in pts]
        lon_span = max(max(lons) - min(lons), 0.002) * (1 + pad)
        lat_span = max(max(lats) - min(lats), 0.002) * (1 + pad)
        cos_lat  = _math.cos(_math.radians(center_lat))
        z_lon = _math.log2(w / 256 * 360 / lon_span)
        z_lat = _math.log2(h / 256 * 180 / (lat_span / cos_lat))
        return max(11, min(16, round(min(z_lon, z_lat), 1)))

    tooltip = {
        "html": (
            "<div style='font-family:sans-serif;padding:6px'>"
            "<b style='font-size:13px'>{marker}. {property}</b>"
            "<br/><span style='color:#aaa;font-size:11px'>{address}</span>"
            "</div>"
        ),
        "style": {
            "background": "rgba(20,20,20,0.88)",
            "color": "white",
            "border-radius": "6px",
        },
    }

    # Per-pin colour. An explicit "color" on the comp ("red"/"navy") wins; else the
    # default (red when the subject is hidden, navy otherwise). Matches the PNG.
    _RED, _NAVY = [210, 40, 40, 240], [25, 90, 200, 230]
    _default_rgba = _RED if not _show_subject else _NAVY
    for _c in visible:
        _col = str(_c.get("color") or "").lower()
        _c["_rgba"] = (_RED if _col in ("red", "subject")
                       else _NAVY if _col in ("navy", "blue", "comp")
                       else _default_rgba)
    comp_layer = pdk.Layer(
        "ScatterplotLayer", data=visible, get_position=["lon", "lat"],
        get_color="_rgba", get_radius=160, radius_min_pixels=10,
        radius_max_pixels=26, pickable=True, auto_highlight=True,
    )
    subj_layer = pdk.Layer(
        "ScatterplotLayer", data=subject_pt, get_position=["lon", "lat"],
        get_color=[210, 40, 40, 240], get_radius=180, radius_min_pixels=11,
        radius_max_pixels=28, pickable=True, auto_highlight=True,
    )
    label_layer = pdk.Layer(
        "TextLayer", data=visible, get_position=["lon", "lat"], get_text="marker",
        get_size=14, get_color=[255, 255, 255, 255],
        get_alignment_baseline="'center'", pickable=False,
    )
    view = pdk.ViewState(longitude=center_lon, latitude=center_lat,
                         zoom=_fit_zoom(all_pts), pitch=0)

    # Original Mapbox basemap when a token is set (best look); token-free Carto
    # "light" (Carto positron) as fallback. Mapbox tiles need a valid token reaching
    # the browser — may be blank on the cloud/corporate, but works locally.
    if token:
        import os as _os
        _os.environ["MAPBOX_API_KEY"] = token
        raw_style = mb.get("style", "streets-v12")
        map_style = (raw_style if raw_style.startswith("mapbox://")
                     else f"mapbox://styles/mapbox/{raw_style}")
        map_prov  = "mapbox"
    else:
        map_style = "light"
        map_prov  = "carto"

    deck = pdk.Deck(
        layers=[comp_layer] + ([subj_layer] if _show_subject else []) + [label_layer],
        initial_view_state=view, tooltip=tooltip,
        map_provider=map_prov, map_style=map_style,
    )
    st.pydeck_chart(deck, use_container_width=True, height=520)

    # ── Map debug (blank-map diagnosis; DevTools is blocked on corporate) ──────
    print(f"[map] provider={map_prov} style={map_style!r} "
          f"mapbox_token={'set(' + str(len(token)) + ')' if token else 'MISSING'} "
          f"comps={len(visible)} subject={'yes' if _show_subject else 'no'} "
          f"center=({center_lat:.5f},{center_lon:.5f})")
    with st.expander("🔧 Map debug info", expanded=False):
        st.json({
            "geo_file":          geo_path.name,
            "map_provider":      map_prov,
            "map_style":         map_style,
            "mapbox_token_set":  bool(token),
            "mapbox_token_len":  len(token or ""),
            "comps_plotted":     len(visible),
            "subject_shown":     bool(_show_subject),
            "center":            [round(center_lat, 5), round(center_lon, 5)],
            "first_comp":        (visible[0] if visible else None),
        })

    # ── Map credit ───────────────────────────────────────────────────────────
    _ss_credit = {}
    try:
        _ss_credit = json.loads((ROOT / "configs" / "shared_settings.json").read_text(encoding="utf-8"))
    except Exception:
        pass
    _geo_prov = _ss_credit.get("geocoding_provider", "mapbox")
    _geo_label = {"mapbox": "Mapbox", "onemap": "OneMap (Singapore)", "google": "Google Maps", "kakao": "Kakao Maps (Korea)"}.get(_geo_prov, _geo_prov.title())
    st.caption(f"Geocoding: {_geo_label}  ·  Map: Mapbox")

    # ── Pin management panel ──────────────────────────────────────────────────
    with st.expander("📍 Manage pins", expanded=False):
        st.caption("Delete pin only removes it from the map. "
                   "Delete pin + row also removes the row from the Excel file. "
                   "Colour toggles a pin between navy and red.")
        _default_color = "red" if not _show_subject else "navy"
        for c in all_comps:
            marker     = str(c["map_marker"])
            is_hidden  = c.get("hidden", False)
            prop_label = f"**{marker}.** {c.get('property', '')}"
            if c.get("address"):
                prop_label += f"  ·  *{c['address']}*"

            col_lbl, col_clr, col_a, col_b = st.columns([5, 2, 2, 2])
            col_lbl.markdown(
                f"~~{prop_label}~~" if is_hidden else prop_label
            )

            if is_hidden:
                if col_a.button("↩️ Restore", key=f"restore_{context}_{marker}"):
                    _update_geo_hidden(geo_path, marker, False)
                    _regen_map_from_geo(geo_path, plot_subject=_show_subject)
                    st.rerun()
            else:
                # Colour toggle (navy ↔ red). Shows the colour it will switch TO.
                _cur  = str(c.get("color") or _default_color).lower()
                _next = "navy" if _cur in ("red", "subject") else "red"
                _lbl  = "🔵 Navy" if _next == "navy" else "🔴 Red"
                if col_clr.button(_lbl, key=f"color_{context}_{marker}",
                                  help="Change this pin's colour"):
                    _update_geo_color(geo_path, marker, _next)
                    _regen_map_from_geo(geo_path, plot_subject=_show_subject)
                    st.rerun()
                if col_a.button("🗑️ Pin only",
                                key=f"del_pin_{context}_{marker}",
                                help="Hide this pin from the map"):
                    _update_geo_hidden(geo_path, marker, True)
                    _regen_map_from_geo(geo_path, plot_subject=_show_subject)
                    st.rerun()
                if col_b.button("🗑️ Pin + row",
                                key=f"del_both_{context}_{marker}",
                                help="Hide pin AND delete this row from the Excel"):
                    _update_geo_hidden(geo_path, marker, True)
                    _delete_excel_row_by_marker(str(excel_path), marker)
                    _regen_map_from_geo(geo_path, plot_subject=_show_subject)
                    st.rerun()

    return True


def _delete_excel_row_by_marker(excel_path: str, marker: str):
    """Remove the row matching map_marker from the Excel preview sheet."""
    import openpyxl as _xl
    try:
        wb = _xl.load_workbook(excel_path)
        ws = wb.active
        # Find the second header row — the comps table header (not the subject one).
        # Scanning only max_row=1 misses it; we must walk all rows and count.
        headers_found = 0
        comp_header_row = None
        col_idx = None
        for row in ws.iter_rows():
            vals = [str(c.value or "").strip() for c in row]
            lower = [v.lower() for v in vals]
            if (any("property" in v for v in lower)
                    and any("marker" in v for v in lower)):
                headers_found += 1
                if headers_found == 2:
                    comp_header_row = row[0].row
                    # Find Map Marker column in this header row
                    for i, v in enumerate(vals):
                        if "marker" in v.lower():
                            col_idx = i + 1   # 1-based
                            break
                    break
        if comp_header_row is None or col_idx is None:
            return
        rows_to_delete = [
            r for r in range(comp_header_row + 1, ws.max_row + 1)
            if str(ws.cell(row=r, column=col_idx).value or "").strip() == marker
        ]
        for r in reversed(rows_to_delete):
            ws.delete_rows(r)
        wb.save(excel_path)
    except Exception as e:
        st.warning(f"Could not delete row from Excel: {e}")


def _update_geo_hidden(geo_path, marker: str, hidden_value: bool):
    """
    Set the 'hidden' flag for a single map marker in the geo JSON sidecar.
    This is the canonical way to toggle pin visibility — the geo JSON is the
    single source of truth; pydeck and the PNG both derive from it.
    """
    import json as _json
    geo_path = Path(geo_path)
    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
        for c in geo.get("comps", []):
            if str(c.get("map_marker", "")) == marker:
                c["hidden"] = hidden_value
        geo_path.write_text(_json.dumps(geo, indent=2), encoding="utf-8")
    except Exception as e:
        st.warning(f"Could not update geo sidecar: {e}")


def _update_geo_color(geo_path, marker: str, color: str):
    """
    Set the per-pin 'color' ("red"/"navy") for a single map marker in the geo JSON
    sidecar. Both the interactive map and the static PNG read this override.
    """
    import json as _json
    geo_path = Path(geo_path)
    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
        for c in geo.get("comps", []):
            if str(c.get("map_marker", "")) == marker:
                c["color"] = color
        geo_path.write_text(_json.dumps(geo, indent=2), encoding="utf-8")
    except Exception as e:
        st.warning(f"Could not update geo sidecar: {e}")


def _regen_map_from_geo(geo_path, plot_subject: bool = True):
    """
    Regenerate the static map PNG from the geo JSON sidecar, plotting only
    comps where hidden=false.  Called immediately after any pin toggle so the
    downloadable PNG stays in sync with the interactive pydeck view.
    plot_subject=False omits the subject star from the saved PNG.
    Silently skips if the sidecar is missing or no Mapbox token is set.
    """
    import json as _json
    from pathlib import Path as _Path
    geo_path = _Path(geo_path)
    if not geo_path.exists():
        return

    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
    except Exception:
        return

    subj  = geo.get("subject", {})
    s_lon = subj.get("lon")
    s_lat = subj.get("lat")
    if s_lon is None or s_lat is None:
        return

    # Only visible (non-hidden) comps with coordinates. 4th element = per-pin
    # colour override ("red"/"navy") from the geo sidecar, if the analyst set one.
    comps_geo = [
        (c["map_marker"], c["lon"], c["lat"], c.get("color"))
        for c in geo.get("comps", [])
        if c.get("lon") is not None and not c.get("hidden", False)
    ]

    mb  = geo.get("mapbox", {})
    tok = _get_mapbox_token()   # always from shared_settings.json
    if not tok:
        return

    map_path = str(geo_path).replace("_geo.json", "_map.png")
    try:
        sys.path.insert(0, str(ROOT / "backend"))
        from generate_sales_comps_map import render_map
        render_map(
            subject_lonlat = (s_lon, s_lat),
            comps          = comps_geo,
            token          = tok,
            output_path    = map_path,
            style          = mb.get("style",   "streets-v12"),
            width          = mb.get("width",   1200),
            height         = mb.get("height",  900),
            padding        = mb.get("padding", 100),
            pin_size       = mb.get("pin_size", "l"),
            plot_subject   = plot_subject,
        )
    except Exception as e:
        st.warning(f"Map regeneration skipped: {e}")


def _regen_map_from_table(excel_path, edited_df):
    """
    After a table save, sync the geo JSON hidden flags with the markers that
    remain in the edited table, then regenerate the map PNG.
    Silently skips if no geo sidecar exists (rent / land comps, old runs).
    """
    import json as _json
    from pathlib import Path as _Path
    excel_path = _Path(excel_path)
    geo_path   = excel_path.parent / excel_path.name.replace(".xlsx", "_geo.json")
    if not geo_path.exists():
        return

    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
    except Exception:
        return

    # Markers still present in the edited table
    remaining_markers = set()
    if "Map Marker" in edited_df.columns:
        remaining_markers = {
            str(v).strip()
            for v in edited_df["Map Marker"].dropna()
            if str(v).strip() not in ("", "★", "*")
        }

    # Sync geo JSON hidden flags: hide any comp whose marker was deleted from table
    if remaining_markers:
        changed = False
        for c in geo.get("comps", []):
            marker = str(c.get("map_marker", ""))
            should_hide = marker not in remaining_markers
            if c.get("hidden", False) != should_hide:
                c["hidden"] = should_hide
                changed = True
        if changed:
            try:
                geo_path.write_text(_json.dumps(geo, indent=2), encoding="utf-8")
            except Exception:
                pass

    # Regenerate PNG from the updated geo JSON
    _regen_map_from_geo(geo_path)


def _geocode_new_rows_and_regen_map(config_path: str, excel_path, edited_df):
    """
    Geocode any rows in edited_df that have no Map Marker (manually added rows),
    append them to the geo JSON sidecar, and regenerate the map.

    Rows that already have a Map Marker are already in the sidecar — skipped.
    """
    import json as _json
    import sys as _sys
    from pathlib import Path as _Path

    excel_path = _Path(excel_path)
    geo_path   = excel_path.parent / excel_path.name.replace(".xlsx", "_geo.json")
    if not geo_path.exists():
        st.warning("No geo sidecar found — run the full analysis first to create the map.")
        return False

    tok = _get_mapbox_token()
    if not tok:
        st.warning("No Mapbox token — cannot geocode new rows.")
        return False

    try:
        geo = _json.loads(geo_path.read_text(encoding="utf-8"))
    except Exception as e:
        st.warning(f"Could not read geo sidecar: {e}")
        return False

    subj_lon = geo.get("subject", {}).get("lon")
    subj_lat = geo.get("subject", {}).get("lat")
    if subj_lon is None or subj_lat is None:
        st.warning("Subject coordinates missing from geo sidecar.")
        return False

    # Country name from deal config for geocoding suffix
    cfg          = load_config(config_path)
    country_name = cfg.get("subject_property", {}).get("country_name", "")
    country_code = cfg.get("country_code", "")
    suffix       = f", {country_name}" if country_name else ""

    # Existing markers in the sidecar
    existing_markers = {str(c.get("map_marker", "")) for c in geo.get("comps", [])}
    # Next available marker number
    used_nums = {int(m) for m in existing_markers if m.isdigit()}
    next_num  = max(used_nums, default=0) + 1

    # Find new rows: those with a blank / "—" Map Marker
    _marker_col   = next((c for c in edited_df.columns if "marker" in c.lower()), None)
    _property_col = next((c for c in edited_df.columns if "property" in c.lower()), None)
    if _property_col is None:
        st.warning("No Property column found in the table.")
        return False

    new_rows = edited_df[
        edited_df[_marker_col].astype(str).str.strip().isin(["", "—", "None"])
        if _marker_col else [True] * len(edited_df)
    ]

    if new_rows.empty:
        st.info("No new rows to geocode — all rows already have map markers.")
        return False

    # Geocode each new row
    _sys.path.insert(0, str(ROOT / "backend"))
    try:
        from generate_comps_map_base import geocode_any as geocode_with_fallbacks
    except ImportError as e:
        st.warning(f"Backend import failed: {e}")
        return False

    geocoded = 0
    failed   = []

    for _, row in new_rows.iterrows():
        name  = str(row.get(_property_col, "")).strip()
        if not name or name in ("—", "None", ""):
            continue
        query = f"{name}{suffix}" if suffix.lower() not in name.lower() else name
        try:
            lon, lat = geocode_with_fallbacks([query], tok, country_code)
            marker   = str(next_num)
            next_num += 1
            geo["comps"].append({
                "map_marker": marker,
                "property":   name,
                "address":    "",
                "lon":        lon,
                "lat":        lat,
                "hidden":     False,
            })
            geocoded += 1
        except Exception as exc:
            failed.append(f"{name}: {exc}")

    if geocoded == 0:
        msgs = "; ".join(failed[:3])
        st.warning(f"Could not geocode any new rows. {msgs}")
        return False

    # Save updated geo sidecar
    geo_path.write_text(_json.dumps(geo, indent=2), encoding="utf-8")

    if failed:
        st.warning(f"{len(failed)} row(s) could not be geocoded: "
                   + "; ".join(failed[:3]))

    # Regenerate PNG
    _regen_map_from_geo(geo_path)
    st.success(f"✅  {geocoded} new row(s) geocoded and added to the map.")
    return True


def _show_results(config_path: str, prefix: str, context: str = "",
                  script: str = ""):
    """Show Excel preview (editable) + download and map preview + download.

    ``context`` must be unique per call site (e.g. "upload", "db", "online")
    so Streamlit generates distinct element IDs for each download button.
    ``script`` is the backend scan script name (e.g. "scan_input_sales_comps.py")
    — used by the records editor re-run and the Refine panel.
    """
    cfg     = load_config(config_path)
    out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
    excels  = sorted(f for f in out_dir.glob(f"{prefix}*.xlsx")
                     if not f.name.startswith("~"))   # ignore lock files
    maps    = sorted(out_dir.glob(f"{prefix}*_map.png"))

    if excels:
        latest = excels[-1]
        _records_path = latest.parent / latest.name.replace(".xlsx", "_records.json")

        # ── Editable Excel preview ─────────────────────────────────────────
        preview_df = _read_excel_preview(str(latest))
        if preview_df is not None and not preview_df.empty:
            st.markdown("#### 📊 Preview")

            # Editable Address column — the analyst fills it when a property name
            # is too rough to geocode. On re-run it takes priority over the name.
            # Pre-fill from any address already in _records.json (matched by marker).
            if not any(str(c).strip().lower() == "address" for c in preview_df.columns):
                _addr_by_marker = {}
                _rp_addr = latest.parent / latest.name.replace(".xlsx", "_records.json")
                try:
                    if _rp_addr.exists():
                        for _r in json.loads(_rp_addr.read_text(encoding="utf-8")):
                            _m = str(_r.get("map_marker", "")).strip()
                            if _m:
                                _addr_by_marker[_m] = str(_r.get("address", "") or "").strip()
                except Exception:
                    pass
                _mk_c = next((c for c in preview_df.columns
                              if "marker" in str(c).lower()), None)
                if _mk_c:
                    preview_df["Address"] = (preview_df[_mk_c].astype(str).str.strip()
                                             .map(lambda m: _addr_by_marker.get(m, "")))
                else:
                    preview_df["Address"] = ""

            # ── Column management ──────────────────────────────────────────
            # _vcols_key — registry of ALL available columns (never overwritten by multiselect)
            # _cvis_key  — user's visibility selection (owned by the multiselect widget)
            # _xcols_key — user-added extra columns not in the Excel schema
            _vcols_key = f"vcols_{context}_{latest.stem}"
            _cvis_key  = f"cvis_{context}_{latest.stem}"
            _xcols_key = f"xcols_{context}_{latest.stem}"
            _rname_key = f"rname_{context}_{latest.stem}"
            if _vcols_key not in st.session_state:
                st.session_state[_vcols_key] = list(dict.fromkeys(preview_df.columns))
            else:
                # Merge any columns the Excel gained since this session started
                _cached = st.session_state[_vcols_key]
                _new = [c for c in preview_df.columns if c not in _cached]
                if _new:
                    st.session_state[_vcols_key] = _cached + _new
                    # Also make new columns visible by default
                    if _cvis_key in st.session_state:
                        st.session_state[_cvis_key] = st.session_state[_cvis_key] + _new
            if _xcols_key not in st.session_state:
                st.session_state[_xcols_key] = []
            if _cvis_key not in st.session_state:
                st.session_state[_cvis_key] = list(st.session_state[_vcols_key])
            if _rname_key not in st.session_state:
                st.session_state[_rname_key] = {}

            with st.expander("🔧  Manage columns"):
                all_cols = list(preview_df.columns) + st.session_state[_xcols_key]
                # _cvis_key is owned by this multiselect — no manual sync needed
                st.multiselect(
                    "Visible columns (drag to reorder)",
                    options=all_cols,
                    key=_cvis_key,
                )

                st.divider()
                ca, cb = st.columns([4, 1])
                new_col = ca.text_input("New column name",
                                        key=f"new_col_{context}_{latest.stem}",
                                        placeholder="e.g. Comments")
                if cb.button("➕ Add", key=f"add_col_{context}_{latest.stem}"):
                    if new_col.strip() and new_col.strip() not in all_cols:
                        st.session_state[_xcols_key].append(new_col.strip())
                        st.session_state[_cvis_key].append(new_col.strip())
                        st.rerun()
                    elif new_col.strip() in all_cols:
                        st.warning(f"Column '{new_col.strip()}' already exists.")

                st.divider()
                _cvis_now = st.session_state.get(_cvis_key, [])
                if _cvis_now:
                    cr, cs, ct = st.columns([2, 2, 1])
                    col_to_rename = cr.selectbox("Rename column", options=_cvis_now,
                                                 key=f"ren_sel_{context}_{latest.stem}")
                    new_name = cs.text_input("New name",
                                             key=f"ren_val_{context}_{latest.stem}",
                                             placeholder="New column name")
                    if ct.button("✏️ Rename", key=f"ren_btn_{context}_{latest.stem}"):
                        if new_name.strip() and new_name.strip() not in all_cols:
                            st.session_state[_rname_key][col_to_rename] = new_name.strip()
                            if col_to_rename in st.session_state[_cvis_key]:
                                idx = st.session_state[_cvis_key].index(col_to_rename)
                                st.session_state[_cvis_key][idx] = new_name.strip()
                            if col_to_rename in st.session_state[_xcols_key]:
                                xi = st.session_state[_xcols_key].index(col_to_rename)
                                st.session_state[_xcols_key][xi] = new_name.strip()
                            st.rerun()
                        elif new_name.strip() in all_cols:
                            st.warning(f"Column '{new_name.strip()}' already exists.")

            # Build display DataFrame
            disp_df = preview_df.rename(columns=st.session_state[_rname_key])
            for xc in st.session_state[_xcols_key]:
                if xc not in disp_df.columns:
                    disp_df[xc] = "—"
            # Show all columns from the current Excel output.
            show_cols = list(disp_df.columns)

            # Strip the Average row — it's a formula row, not editable data.
            # New rows added by the user will naturally appear above it.
            # _save_edited_preview re-writes the Average row after saving comps.
            disp_df = disp_df[~disp_df.apply(
                lambda r: any(str(v).strip().lower() == "average" for v in r),
                axis=1
            )].reset_index(drop=True)

            st.caption("Edit cells directly · **＋** adds a row · select row + **Delete** removes it")
            # Display Address right after the Property column. Kept LAST in the
            # dataframe (so the position-based Excel writer still aligns); this only
            # changes display order, not the underlying column order.
            _col_order = list(disp_df.columns)
            _addr_c = next((c for c in _col_order
                            if str(c).strip().lower() == "address"), None)
            if _addr_c:
                _col_order.remove(_addr_c)
                _pi = next((i for i, c in enumerate(_col_order)
                            if "property" in str(c).lower()), None)
                _col_order.insert(_pi + 1 if _pi is not None else len(_col_order), _addr_c)
            edited_df = st.data_editor(
                disp_df,
                num_rows="dynamic",
                hide_index=True,
                use_container_width=True,
                height=min(500, 60 + len(disp_df) * 38),
                column_order=_col_order,
                key=f"preview_editor_{context}_{latest.stem}",
            )

            _save_col, _rerun_col2 = st.columns([1, 1])
            _do_quick_save  = _save_col.button(
                "💾  Save Edits & Update",
                key=f"save_edit_{context}_{latest.stem}",
                help="Save display changes (name, address, date). Does not recalculate metrics.",
            )
            _do_rerun_save  = _rerun_col2.button(
                "🔄  Re-Run & Update",
                key=f"save_rerun_{context}_{latest.stem}",
                help="Save all changes and re-run the pipeline. Use when price, GFA, or tenure was edited.",
            )
            if _do_quick_save or _do_rerun_save:
                rev_rename = {v: k for k, v in st.session_state[_rname_key].items()}
                save_df = edited_df.rename(columns=rev_rename)
                # Renumber map markers sequentially before saving
                _mk_col = next((c for c in save_df.columns
                                if "marker" in c.lower()), None)
                _pr_col = next((c for c in save_df.columns
                                if "property" in c.lower()), None)
                _marker_map = {}
                _sync_df = save_df             # original markers — for records sync
                if _mk_col:
                    _is_subj = save_df[_mk_col].astype(str).str.strip() == "★"
                    _comp_mask = ~_is_subj
                    _old_markers = save_df.loc[_comp_mask, _mk_col].astype(str).tolist()
                    _new_markers = [str(i + 1) for i in range(_comp_mask.sum())]
                    _marker_map  = dict(zip(_old_markers, _new_markers))
                    _sync_df = save_df.copy()  # snapshot with ORIGINAL markers
                    save_df = save_df.copy()
                    save_df.loc[_comp_mask, _mk_col] = _new_markers

                if _save_edited_preview(str(latest), save_df):
                    # Rebuild geo.json with renumbered markers
                    import json as _json
                    _geo_path = latest.parent / latest.name.replace(".xlsx", "_geo.json")
                    if _geo_path.exists() and _mk_col:
                        try:
                            _geo = _json.loads(_geo_path.read_text(encoding="utf-8"))
                            _geo_by_marker = {
                                str(c.get("map_marker", "")): c
                                for c in _geo.get("comps", [])
                            }
                            _rev_marker_map = {
                                new: old for old, new in _marker_map.items()
                            }
                            _new_comps = []
                            for _, _row in save_df.iterrows():
                                _m = str(_row.get(_mk_col, "")).strip()
                                if not _m or _m == "★":
                                    continue
                                _p      = str(_row.get(_pr_col, "")).strip() if _pr_col else ""
                                _old_m  = _rev_marker_map.get(_m, _m)
                                _entry  = dict(_geo_by_marker.get(_old_m, {}))
                                _entry["map_marker"] = _m
                                _entry["property"]   = _p
                                _new_comps.append(_entry)
                            _geo["comps"] = _new_comps

                            # Geocode any new rows that have no coordinates yet
                            _needs_geo = [c for c in _geo["comps"]
                                          if c.get("lon") is None and c.get("property")]
                            if _needs_geo:
                                _tok = _get_mapbox_token()
                                if _tok:
                                    _cfg          = load_config(config_path)
                                    _country_name = _cfg.get("subject_property", {}).get("country_name", "")
                                    _country_code = _cfg.get("country_code", "")
                                    _suffix       = f", {_country_name}" if _country_name else ""
                                    import sys as _sys
                                    _sys.path.insert(0, str(ROOT / "backend"))
                                    try:
                                        from generate_comps_map_base import geocode_any as _gwf
                                        _n_ok, _n_fail = 0, []
                                        for _c in _needs_geo:
                                            _name = _c["property"]
                                            _q    = (f"{_name}{_suffix}"
                                                     if _suffix.lower() not in _name.lower()
                                                     else _name)
                                            try:
                                                _c["lon"], _c["lat"], _ = _gwf(
                                                    [_q], _tok, _country_code)
                                                _n_ok += 1
                                            except Exception as _exc:
                                                _n_fail.append(f"{_name}: {_exc}")
                                        if _n_ok:
                                            st.info(f"📍 {_n_ok} new row(s) geocoded and added to map.")
                                        if _n_fail:
                                            st.warning("Could not geocode: "
                                                       + "; ".join(_n_fail[:3]))
                                    except ImportError as _ie:
                                        st.warning(f"Geocoding skipped: {_ie}")
                                else:
                                    st.warning("No Mapbox token — new rows saved to Excel "
                                               "but not added to map.")

                            _geo_path.write_text(
                                _json.dumps(_geo, indent=2), encoding="utf-8")
                        except Exception as _ge:
                            st.warning(f"Could not update geo sidecar: {_ge}")

                    # Always sync _records.json so deletions/edits persist.
                    # Pass the ORIGINAL-marker snapshot + remap so removals match
                    # _records.json (then survivors are renumbered to match).
                    _sync_records_json(str(_records_path), _sync_df, _marker_map)
                    if _do_rerun_save and script and _records_path.exists():
                        st.info("Re-running analysis with updated records …")
                        _rerun_ok = _run_script(
                            script, config_path,
                            ["--map", "--from-records", str(_records_path)],
                        )
                        if _rerun_ok:
                            st.success("✅  Re-run complete — table and metrics updated.")
                        else:
                            st.warning("Re-run failed — check the error log above.")
                    else:
                        st.success("✅  Changes saved to Excel file.")
                    _regen_map_from_table(latest, save_df)
                    st.rerun()


        # ── Geocoding Status panel ────────────────────────────────────────────
        if _records_path.exists():
            try:
                _geo_records = json.loads(
                    _records_path.read_text(encoding="utf-8"))
                _CONFIDENT_PROVIDERS = {"onemap", "kakao", "google"}
                _confident = [r for r in _geo_records
                              if r.get("_geo_provider") in _CONFIDENT_PROVIDERS]
                _fallback  = [r for r in _geo_records
                              if str(r.get("_geo_provider") or "").startswith("mapbox")]
                _failed    = [r for r in _geo_records
                              if r.get("_geo_provider") == "failed"]
                _unknown   = [r for r in _geo_records
                              if "_geo_provider" not in r]

                # Label confident metric by active provider
                _local_ss = {}
                try:
                    _local_ss = json.loads(
                        (ROOT / "configs" / "shared_settings.json").read_text(encoding="utf-8"))
                except Exception:
                    pass
                _active_geo = _local_ss.get("geocoding_provider", "mapbox").lower()
                _confident_label = {
                    "onemap": "✅ Confident (OneMap)",
                    "kakao":  "✅ Confident (Kakao)",
                    "google": "✅ Confident (Google)",
                }.get(_active_geo, "✅ Confident")

                _has_issues = bool(_fallback or _failed)
                _geo_label  = (
                    f"📍 Geocoding Status — "
                    f"✅ {len(_confident)}  "
                    f"⚠️ {len(_fallback)}  "
                    f"❌ {len(_failed)}"
                    + ("  _(flag for review)_" if _has_issues else "")
                )
                with st.expander(_geo_label, expanded=_has_issues):
                    g1, g2, g3 = st.columns(3)
                    g1.metric(_confident_label,      len(_confident))
                    g2.metric("⚠️ Mapbox fallback",  len(_fallback))
                    g3.metric("❌ Not plotted",       len(_failed))

                    def _prop_name(r):
                        return (r.get("property_name") or r.get("property")
                                or r.get("site_name") or "—")

                    if _fallback:
                        st.write("")
                        st.markdown(
                            "**⚠️ Mapped via Mapbox — position may be approximate:**")
                        st.dataframe(
                            pd.DataFrame([
                                {"Property": _prop_name(r),
                                 "Reason":   r.get("_geo_note", "")}
                                for r in _fallback
                            ]),
                            hide_index=True, use_container_width=True,
                        )

                    if _failed:
                        st.write("")
                        st.markdown(
                            "**❌ Not plotted — manual address correction needed:**")
                        st.dataframe(
                            pd.DataFrame([
                                {"Property": _prop_name(r),
                                 "Error":    r.get("_geo_note", "")}
                                for r in _failed
                            ]),
                            hide_index=True, use_container_width=True,
                        )
                        st.caption(
                            "Correct the property name or address in the Preview "
                            "table above and click **🔄 Re-Run** to retry geocoding.")

                    if _unknown:
                        st.caption(
                            f"ℹ️ {len(_unknown)} record(s) pre-date this feature — "
                            "re-run the pipeline to populate geocoding status.")
            except Exception:
                pass   # never let the status panel break the results page

        # ── Refine panel (re-run from records with analyst instructions) ──────
        if script and _records_path.exists():
            _refine_log_key     = f"refine_log_{context}_{latest.stem}"
            _refine_done_key    = f"refine_done_{context}_{latest.stem}"
            _refine_expand_key  = f"refine_log_exp_{context}_{latest.stem}"

            # After a successful refinement + rerun, show result above the panel
            if st.session_state.get(_refine_done_key):
                st.success("✅  Refinement applied — table updated above.")
                if st.session_state.get(_refine_log_key):
                    _auto_expand = st.session_state.pop(_refine_expand_key, False)
                    with st.expander("📋  Last refinement log", expanded=_auto_expand):
                        st.code(st.session_state[_refine_log_key], language="text")

            st.write("")
            with st.expander("🔄  Refine This Output", expanded=False):
                st.caption(
                    "Describe the changes you want — the model will filter or "
                    "modify the records list, then re-generate the table.  "
                    "Extraction is skipped, so this typically takes **1–2 min**."
                )
                _refine_feedback = st.text_area(
                    "What to change in the next run",
                    height=110,
                    placeholder=(
                        "e.g. Remove any comps sold before 2022.\n"
                        "Only keep freehold and 999-year leasehold properties.\n"
                        "Exclude the outlier at price > S$400M."
                    ),
                    key=f"comp_refinement_{context}_{latest.stem}",
                )
                _refine_engine_label = st.radio(
                    "Engine",
                    ["Standard", "Advanced (sandbox)"],
                    horizontal=True,
                    key=f"comp_refine_engine_{context}_{latest.stem}",
                    help=(
                        "Standard: fast, safe rule engine — best for filters on price, "
                        "date, distance, category, top-N. Runs on our server.\n\n"
                        "Advanced (sandbox): OpenAI Code Interpreter for complex logic "
                        "(e.g. 'within 20% of the subject's size', 'office OR retail "
                        "except keep anything over S$1bn'). Slower; the comp data is sent "
                        "to OpenAI's sandbox — use for non-confidential deals."
                    ),
                )
                if _refine_engine_label.startswith("Advanced"):
                    st.caption("⚠️  Advanced sends the comp list to OpenAI's sandbox — "
                               "avoid for confidential deals.")
                if st.button(
                    "🔄  Apply Changes & Re-run",
                    type="primary",
                    disabled=not _refine_feedback.strip(),
                    key=f"comp_refine_btn_{context}_{latest.stem}",
                ):
                    # Clear previous run state before starting
                    st.session_state.pop(_refine_done_key, None)
                    st.session_state.pop(_refine_log_key, None)
                    _rf = tempfile.NamedTemporaryFile(
                        mode="w", suffix=".txt", delete=False,
                        dir=ROOT / "configs", encoding="utf-8",
                    )
                    _rf.write(_refine_feedback.strip())
                    _rf.close()
                    _refine_flags = ["--map", "--from-records", str(_records_path),
                                     "--refinement-file", _rf.name]
                    if _refine_engine_label.startswith("Advanced"):
                        _refine_flags += ["--refine-engine", "code_interpreter"]
                    _ok = _run_script(
                        script, config_path,
                        _refine_flags,
                        expand_log=True,
                        log_state_key=_refine_log_key,
                    )
                    try: os.unlink(_rf.name)
                    except: pass
                    if _ok:
                        st.session_state[_refine_done_key] = True
                        st.session_state[_refine_expand_key] = True  # auto-expand log once
                        st.rerun()  # refresh preview immediately

        # ── Download button  (key must be unique across all call sites) ────
        with open(latest, "rb") as fh:
            st.download_button(
                f"⬇️  Download {latest.name}", fh.read(), latest.name,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"dl_excel_{context}_{latest.stem}",
                use_container_width=True)

    # ── Interactive map (from geo sidecar) or static PNG fallback ────────────
    if excels:
        latest     = excels[-1]
        geo_path   = latest.parent / latest.name.replace(".xlsx", "_geo.json")
        st.markdown("#### 🗺️ Location Map")
        shown = _show_interactive_map(geo_path, latest, context)
        if not shown and maps:
            # Fallback to static PNG if no geo sidecar (rent/land comps, old runs)
            st.image(str(maps[-1]), caption=maps[-1].name, use_container_width=True)

    if maps:
        with open(maps[-1], "rb") as fh:
            st.download_button(
                f"⬇️  Download {maps[-1].name}", fh.read(), maps[-1].name,
                "image/png",
                key=f"dl_map_{context}_{maps[-1].stem}",
                use_container_width=True)


def _extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    if name.endswith((".xlsx", ".xls")):
        import openpyxl
        wb    = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                row_txt = "  ".join(str(c.value) for c in row if c.value not in (None, ""))
                if row_txt.strip():
                    lines.append(row_txt)
        return "\n".join(lines)
    if name.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            return "\n".join(p.extract_text() or "" for p in reader.pages[:6])
        except ImportError:
            return "[PDF extraction requires pypdf — run: pip install pypdf]"
    return data.decode("utf-8", errors="ignore")


def _markdown_to_docx(md_text: str) -> bytes:
    """Convert a markdown string to a .docx file, returned as bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
    except ImportError:
        raise ImportError("python-docx required — run: pip install python-docx")

    doc = Document()
    # Narrow margins
    for sec in doc.sections:
        sec.left_margin  = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.top_margin   = Inches(1.0)
        sec.bottom_margin = Inches(1.0)

    _append_markdown_to_doc(doc, md_text)

    _apply_docx_font(doc, "Arial Narrow", 10)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _append_markdown_to_doc(doc, md_text: str):
    """Append markdown (headings, paragraphs, metadata lines) to an existing
    python-docx Document. Shared by _markdown_to_docx and _build_combined_docx."""
    import re as _re
    from docx.shared import Inches
    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped == "---":
            doc.add_paragraph("─" * 60)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            p.paragraph_format.left_indent = Inches(0.4)
        elif stripped.startswith("**") and ":**" in stripped:
            p = doc.add_paragraph()
            m = _re.match(r'\*\*(.+?):\*\*\s*(.*)', stripped)
            if m:
                r = p.add_run(m.group(1) + ": ")
                r.bold = True
                p.add_run(m.group(2))
            else:
                p.add_run(stripped)
        else:
            doc.add_paragraph(stripped)


# ═════════════════════════════════════════════════════════════════════════════
# PGIM styled comp table  +  combined-report export  (used by the Overview)
# ═════════════════════════════════════════════════════════════════════════════

_PGIM_NAVY = "#1F3864"

# Output prefix → (section heading, export section title, subject banner, comps
# banner, average-column keywords). Single source of truth for the three comp
# types across the Overview preview and the combined Word export. Order here is
# the render order (Rent → Asset Sales → Land; Investment Rationale follows).
_COMP_TYPES = [
    ("Rent_Comps",              "Rent Comparables",
     "Rent Comparables",
     "Subject Rents", "Leasing Comparables", ("psf",)),
    ("Transaction_Comparables", "Transaction Comparables (Asset Sales)",
     "Transaction Comparables (Asset Sales)",
     "Subject Sales", "Comparable Asset Sales", ("psf", "cap rate", "capitalisation")),
    ("Land_Sale_Comps",         "Transaction Comparables (Land Sales)",
     "Transaction Comparables (Land Sales)",
     "Subject Land Sales", "Comparable Land Sales", ("psf", "tenure")),
]


def _input_list(v) -> list:
    """Config input value → list of paths (handles str, list, or None)."""
    if not v:
        return []
    return [v] if isinstance(v, str) else list(v)


def _rationale_body(md_text: str) -> str:
    """Return just the rationale prose — strip the metadata header block
    (title / Address / Asset Class / Country / Generated / Market Reports Used /
    the LLM-generated notice) and the '## Input Summary' section. Used by the
    Overview preview and the Word export; the detail subtab shows the full doc."""
    lines = md_text.splitlines()
    # 1) Drop the leading metadata block — everything up to & incl. the first '---'.
    start = 0
    for i, ln in enumerate(lines):
        if ln.strip() == "---":
            start = i + 1
            break
    rest = lines[start:]
    # 2) If an 'Input Summary' section leads, drop it up to & incl. its closing '---'.
    j = 0
    while j < len(rest) and not rest[j].strip():
        j += 1
    if j < len(rest) and rest[j].strip().lower().lstrip("# ").startswith("input summary"):
        k = j + 1
        while k < len(rest) and rest[k].strip() != "---":
            k += 1
        rest = rest[k + 1:]
    return "\n".join(rest).strip()


def _save_uploads_append(cfg: dict, excel_key: str, pdf_key: str, image_key: str,
                         uploaded_files) -> dict:
    """Save uploaded files to Input_files/ and UNION them into the config's input
    lists (append, not replace — deduped by path). Mutates and returns cfg."""
    img_exts = {".png", ".jpg", ".jpeg"}
    buckets = {excel_key: _input_list(cfg.get(excel_key)),
               pdf_key:   _input_list(cfg.get(pdf_key)),
               image_key: _input_list(cfg.get(image_key))}
    for uf in uploaded_files:
        suf = Path(uf.name).suffix.lower()
        key = image_key if suf in img_exts else (pdf_key if suf == ".pdf" else excel_key)
        sp = ROOT / "Input_files" / uf.name
        sp.parent.mkdir(parents=True, exist_ok=True)
        sp.write_bytes(uf.getvalue())
        rel = str(sp.relative_to(ROOT))
        if rel not in buckets[key]:
            buckets[key].append(rel)
    for k, v in buckets.items():
        if v:
            cfg[k] = v[0] if len(v) == 1 else v
    return cfg


def _compute_average(header, comps, keywords):
    """Compute an Average row from the comp rows for columns whose header matches
    any keyword (e.g. 'psf', 'cap rate', 'tenure'). Needed because the Excel's
    =AVERAGEIF formulas have no cached value when read via openpyxl data_only."""
    if not comps or not keywords:
        return None
    avg = [""] * len(header)
    if header:
        avg[0] = "Average"
    any_val = False
    for j, h in enumerate(header):
        if not any(k in h.lower() for k in keywords):
            continue
        nums, had_pct = [], False
        for row in comps:
            if j < len(row):
                raw = str(row[j]).strip()
                if "%" in raw:
                    had_pct = True
                v = raw.replace(",", "").replace("%", "").strip()
                try:
                    nums.append(float(v))
                except ValueError:
                    pass
        if nums:
            m = sum(nums) / len(nums)
            s = f"{m:,.1f}" if abs(m) < 1000 else f"{m:,.0f}"
            avg[j] = (s + "%") if had_pct else s
            any_val = True
    return avg if any_val else None


def _read_pgim_grid(excel_path, avg_keywords=None):
    """Read a formatted comp Excel into a structured PGIM layout (values only —
    colours are applied by row role downstream):
        {"header": [labels], "subject": [[vals], ...],
         "comps":  [[vals], ...], "average": [vals] | None}
    Column order is preserved and the internal 'Source' column is dropped. If the
    Excel's average row is empty (formula with no cached value), an Average row is
    computed from the comps for the columns matching ``avg_keywords``."""
    import openpyxl
    from openpyxl.utils import get_column_letter
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception:
        return None
    ws = wb.active
    hidden = {i for i in range(1, ws.max_column + 1)
              if ws.column_dimensions.get(get_column_letter(i),
                 type("_", (), {"hidden": False})()).hidden}
    rows = list(ws.iter_rows())

    def _vis(row):
        return [c for j, c in enumerate(row, 1) if j not in hidden]

    def _is_avg_fill(cell):
        try:
            rgb = cell.fill.fgColor.rgb
            return isinstance(rgb, str) and rgb.upper().endswith("D6DCE4")
        except Exception:
            return False

    # First column-header row (contains 'property' + 'marker').
    header_idx = None
    for i, row in enumerate(rows):
        strs = [str(c.value or "").lower() for c in _vis(row)]
        if (sum(1 for s in strs if s) >= 3
                and any("property" in s for s in strs)
                and any("marker" in s for s in strs)):
            header_idx = i
            break
    if header_idx is None:
        return None

    vis_header = _vis(rows[header_idx])
    # Hide the internal "Source" column from the PGIM view + Word export.
    drop = {idx for idx, c in enumerate(vis_header)
            if str(c.value or "").strip().lower() == "source"}

    def _keep(cells):
        return [c for idx, c in enumerate(cells) if idx not in drop]

    def _vals(cells):
        return ["" if c.value in (None, "") else str(c.value).replace("\n", " ").strip()
                for c in _keep(cells)]

    def _is_header(cells):
        low = [str(c.value or "").lower() for c in _keep(cells)]
        return any("property" in s for s in low) and any("marker" in s for s in low)

    def _is_blank(cells):
        return all(c.value in (None, "") for c in _keep(cells))

    def _is_average(cells):
        if any(str(c.value or "").strip().lower() == "average" for c in _keep(cells)):
            return True
        return any(_is_avg_fill(c) for c in cells)

    header = [str(c.value or "").replace("\n", " ").strip() for c in _keep(vis_header)]
    subject, comps, average = [], [], None
    phase = 0  # 0 = subject block, 1 = comps block (after the blank / repeat header)
    for row in rows[header_idx + 1:]:
        cells = _vis(row)
        if _is_average(cells):
            _v = _vals(cells)
            if any(_v):                       # keep only a non-empty average row
                average = _v
            phase = 1
            continue
        if _is_blank(cells):
            if phase == 0 and subject:        # blank after the subject block
                phase = 1
            continue
        if _is_header(cells):                 # repeated header before comps
            phase = 1
            continue
        (subject if phase == 0 else comps).append(_vals(cells))
    if not average:                       # Excel formula average → compute here
        average = _compute_average(header, comps, avg_keywords)
    return {"header": header, "subject": subject, "comps": comps, "average": average}


def _fg_for(bg_hex) -> str:
    """White or navy text for a given background hex, chosen for contrast."""
    if not bg_hex:
        return "#1A3A5C"
    try:
        r, g, b = int(bg_hex[1:3], 16), int(bg_hex[3:5], 16), int(bg_hex[5:7], 16)
        return "#ffffff" if (0.299 * r + 0.587 * g + 0.114 * b) < 140 else "#1A3A5C"
    except Exception:
        return "#1A3A5C"


def _pgim_table_html(excel_path, subject_banner: str, comp_banner: str,
                     avg_keywords=None):
    """Formatted comp Excel → static PGIM-standard HTML table:
       navy banner → column names → subject → blank → navy banner → comps →
       grey bold average. None if unreadable."""
    import html as _html
    grid = _read_pgim_grid(excel_path, avg_keywords)
    if not grid or not grid["header"]:
        return None
    header = grid["header"]
    ncol = len(header)
    # Horizontal rules only (no vertical lines). Banners left, other cells centered.
    _cb = "border-bottom:1px solid #cfd8e3;padding:5px 9px;text-align:left;"
    _cd = "border-bottom:1px solid #cfd8e3;padding:5px 9px;text-align:center;"

    def _banner(text):
        return (f'<tr><td colspan="{ncol}" style="{_cb}background:{_PGIM_NAVY};'
                f'color:#ffffff;font-weight:700;">{_html.escape(text)}</td></tr>')

    def _row(vals, bold=False, bg="#ffffff"):
        vals = (list(vals) + [""] * ncol)[:ncol]
        w = "font-weight:700;" if bold else ""
        tds = "".join(f'<td style="{_cd}background:{bg};color:#000000;{w}">'
                      f'{_html.escape(v)}</td>' for v in vals)
        return f"<tr>{tds}</tr>"

    out = ['<div style="overflow-x:auto">',
           '<table style="border-collapse:collapse;width:100%;'
           'font-family:Segoe UI,Arial,sans-serif;font-size:12px;">']
    out.append(_banner(subject_banner))
    out.append(_row(header, bold=True))               # column names — black, no fill
    for s in grid["subject"]:
        out.append(_row(s))
    out.append(f'<tr><td colspan="{ncol}" style="border:0;height:10px;"></td></tr>')
    out.append(_banner(comp_banner))
    for c in grid["comps"]:
        out.append(_row(c))
    if grid["average"]:
        out.append(_row(grid["average"], bold=True, bg="#D6DCE4"))
    out.append("</table></div>")
    return "".join(out)


def _set_cell_bg(cell, hex_no_hash: str):
    """Shade a python-docx table cell (adds a w:shd element)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_no_hash)
    tcPr.append(shd)


def _table_horizontal_borders(table, color: str = "808080", sz: str = "4"):
    """Horizontal rules only — remove all vertical (inner + outer) lines."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none"); el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _add_pgim_table_to_doc(doc, excel_path, subject_banner: str, comp_banner: str,
                           avg_keywords=None) -> bool:
    """Add a native PGIM-standard Word table: navy banner → column names →
    subject → blank → navy banner → comps → grey bold average. Horizontal rules
    only (no vertical lines); every cell centered except the navy banners."""
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    grid = _read_pgim_grid(excel_path, avg_keywords)
    if not grid or not grid["header"]:
        return False
    header = grid["header"]
    ncol = len(header)

    # Ordered row plan: (kind, values)
    seq = [("banner", subject_banner), ("header", header)]
    seq += [("data", s) for s in grid["subject"]]
    seq.append(("blank", None))
    seq.append(("banner", comp_banner))
    seq += [("data", c) for c in grid["comps"]]
    if grid["average"]:
        seq.append(("average", grid["average"]))

    table = doc.add_table(rows=len(seq), cols=ncol)
    _table_horizontal_borders(table)

    for ridx, (kind, vals) in enumerate(seq):
        row = table.rows[ridx]
        if kind == "banner":
            merged = row.cells[0]
            for j in range(1, ncol):
                merged = merged.merge(row.cells[j])
            merged.text = vals
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_bg(merged, _PGIM_NAVY.lstrip("#"))
            for p in merged.paragraphs:               # banner stays left-aligned
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            continue
        if kind == "blank":
            continue                                   # empty separator row
        cvals = (list(vals) + [""] * ncol)[:ncol]
        grey = kind == "average"
        bold = kind in ("header", "average")
        for j, v in enumerate(cvals):
            cell = row.cells[j]
            cell.text = v
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if grey:
                _set_cell_bg(cell, "D6DCE4")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    if bold:
                        r.font.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return True


def _apply_docx_font(doc, name: str = "Arial Narrow", size_pt: float = 10):
    """Force one font + size across the whole document — the Normal/heading
    styles plus every run in body paragraphs and table cells. Colour and bold
    already set on runs are preserved (only name + size are changed)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            stl = doc.styles[style_name]
            stl.font.name = name
            stl.font.size = Pt(size_pt)
            stl.font.italic = False
        except Exception:
            pass

    def _set_runs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                r.font.name = name
                r.font.size = Pt(size_pt)
                r.font.italic = False       # kill any inherited italic
                # Pin the East-Asian / complex-script slots too, so Word doesn't
                # substitute a different face for any character.
                try:
                    rpr = r._element.get_or_add_rPr()
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = rpr.makeelement(qn("w:rFonts"), {})
                        rpr.append(rfonts)
                    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                        rfonts.set(qn(attr), name)
                except Exception:
                    pass

    _set_runs(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:      # tight rows — no para spacing
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                _set_runs(cell.paragraphs)


def _png_size(path):
    """(width, height) in pixels from a PNG header, else (None, None)."""
    try:
        with open(path, "rb") as f:
            head = f.read(24)
        if len(head) >= 24 and head[:8] == b"\x89PNG\r\n\x1a\n":
            import struct
            return struct.unpack(">II", head[16:24])
    except Exception:
        pass
    return None, None


def _add_map_fit(doc, map_path, usable_w, usable_h):
    """Insert a map scaled to the largest size that fits the content box
    (usable_w × usable_h), preserving aspect ratio."""
    _pw, _ph = _png_size(str(map_path))
    try:
        if _pw and _ph and (_pw / _ph) < (usable_w / usable_h):
            doc.add_picture(str(map_path), height=usable_h)   # tall image → cap height
        else:
            doc.add_picture(str(map_path), width=usable_w)    # wide image → fill width
    except Exception:
        try:
            doc.add_picture(str(map_path), width=usable_w)
        except Exception:
            pass


def _build_combined_docx(deal_name: str, cfg: dict):
    """One Word file: Sales / Land / Rent PGIM tables + location maps, then the
    Investment Rationale. Returns docx bytes, or None if there is nothing yet."""
    from docx import Document
    from docx.shared import Inches, Cm
    from docx.enum.section import WD_ORIENT
    out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
    subj = cfg.get("subject_property", {})

    doc = Document()
    for sec in doc.sections:
        # US Letter, landscape (27.94 cm × 21.59 cm). python-docx needs the
        # width/height swapped manually when setting orientation.
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width  = Cm(27.94)
        sec.page_height = Cm(21.59)
        sec.left_margin = sec.right_margin = Inches(0.5)
        sec.top_margin  = sec.bottom_margin = Inches(0.5)
    _sec0 = doc.sections[0]
    _usable_w = _sec0.page_width - _sec0.left_margin - _sec0.right_margin
    _usable_h = _sec0.page_height - _sec0.top_margin - _sec0.bottom_margin
    doc.add_heading(subj.get("property_name", deal_name), level=0)
    if subj.get("address"):
        doc.add_paragraph(subj["address"])

    any_content = False
    for prefix, _heading, title, _sub_banner, _comp_banner, _avg_kw in _COMP_TYPES:
        xls, used = _latest_comp_excel(out_dir, prefix)
        if not xls:
            continue
        any_content = True
        doc.add_heading(title, level=1)
        _add_pgim_table_to_doc(doc, str(xls), _sub_banner, _comp_banner, _avg_kw)
        maps = sorted(out_dir.glob(f"{used}*_map.png"))
        if maps:
            _add_map_fit(doc, maps[-1], _usable_w, _usable_h)

    rat = out_dir / "Investment_Rationale.md"
    if rat.exists():
        any_content = True
        doc.add_page_break()
        doc.add_heading("Investment Rationale", level=1)
        try:
            _append_markdown_to_doc(doc, _rationale_body(rat.read_text(encoding="utf-8")))
        except Exception:
            pass

    if not any_content:
        return None
    _apply_docx_font(doc, "Arial Narrow", 10)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _get_mapbox_token() -> str:
    """
    Read the Mapbox token from shared_settings.json — the single source of truth.
    Never read from individual deal configs or geo JSON files.
    """
    ss_path = ROOT / "configs" / "shared_settings.json"
    try:
        return json.loads(ss_path.read_text(encoding="utf-8")).get("mapbox_token", "")
    except Exception:
        return ""


def _load_shared_settings_to_env() -> None:
    """Load API keys from shared_settings.json into os.environ so subprocesses inherit them."""
    ss_path = ROOT / "configs" / "shared_settings.json"
    try:
        ss = json.loads(ss_path.read_text(encoding="utf-8"))
        if ss.get("openai_api_key") and not os.environ.get("OPENAI_API_KEY"):
            os.environ["OPENAI_API_KEY"] = ss["openai_api_key"]
    except Exception:
        pass


_load_shared_settings_to_env()


def _ollama_base_url() -> str:
    """
    Read the Ollama base URL from any available deal config.
    Model selection is always driven by the sidebar — never read from config.
    """
    for f in sorted((ROOT / "configs").glob("deal_config*.json")):
        try:
            url = json.loads(f.read_text(encoding="utf-8")).get("llm", {}).get("ollama", {}).get("base_url")
            if url:
                return url
        except Exception:
            pass
    return "http://localhost:11434"


# ── Preview table helpers ─────────────────────────────────────────────────────
# _PREVIEW_ROWS defines every config field that is shown in the New Deal review
# table (Step 2).  Each entry is a 3-tuple:
#   (section_label, display_label, config_key)
# section_label — groups fields under "Property", "Deal", or "Market" in the UI
# display_label — human-readable name shown in the Field column
# config_key    — the key name used in deal_config_*.json subject_property dict
_PREVIEW_ROWS = [
    ("Property", "Deal Name",                 "deal_name"),
    ("Property", "Property Name",             "property_name"),
    ("Property", "Address",                   "address"),
    ("Property", "Asset Class",               "asset_class"),
    ("Property", "GFA",                       "gfa_sf"),
    ("Property", "Quality",                   "quality"),
    ("Property", "Asset Type / Structure",    "asset_type"),
    ("Deal",     "Sale Date Label",           "sale_date"),
    ("Deal",     "Remaining Leasehold (yrs)", "remaining_leasehold_yrs"),
    ("Deal",     "Price (millions)",          "price_sgd_m"),
    ("Deal",     "Price Unit (T / M / B)",     "price_unit"),
    ("Deal",     "Area Unit (sqm or SF)",     "area_unit"),
    ("Deal",     "Cap Rate (decimal)",        "ftm_noi_cap_rate"),
    ("Market",   "Country",                   "country_name"),
    ("Market",   "Country Code",              "country_code"),
    ("Market",   "Currency",                  "currency"),
    ("Market",   "Currency Symbol",           "currency_symbol"),
    ("Market",   "GFA Unit",                  "gfa_unit"),
    ("Market",   "Land Zoning",               "land_zoning"),
    ("Market",   "Location Descriptor",       "location"),
    ("Market",   "Asset Search Keyword",      "asset_search_keyword"),
    ("Market",   "Submarket Keywords (csv)",  "submarket_keywords"),
    ("Market",   "Broader Market Query",      "broader_market_query"),
]


def _fields_to_df(fields: dict) -> pd.DataFrame:
    rows = []
    for section, label, key in _PREVIEW_ROWS:
        val = fields.get(key, "")
        if isinstance(val, list):
            val = ", ".join(str(v) for v in val)
        rows.append({"Section": section, "Field": label,
                     "_key": key, "Value": str(val) if val is not None else ""})
    return pd.DataFrame(rows)


def _df_to_fields(df: pd.DataFrame) -> dict:
    fields = {}
    for _, row in df.iterrows():
        key = row["_key"]
        val = str(row["Value"]).strip()
        if key == "submarket_keywords":
            fields[key] = [k.strip() for k in val.split(",") if k.strip()]
        elif key in ("gfa_sf", "remaining_leasehold_yrs"):
            try:    fields[key] = int(float(val))
            except: fields[key] = 0
        elif key in ("price_sgd_m", "ftm_noi_cap_rate"):
            try:    fields[key] = float(val) if val else None
            except: fields[key] = None
        else:
            fields[key] = val
    return fields


# ── Session state defaults ────────────────────────────────────────────────────
# Initialise keys that must exist before the first render pass.
# Streamlit re-runs the entire script on every user interaction, so these guards
# ensure state is only set on the very first load and not reset on each re-run.
for _k, _v in {"nd_step": 1, "nd_preview_fields": None}.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ═════════════════════════════════════════════════════════════════════════════
# SIDEBAR  — Navigation
# ═════════════════════════════════════════════════════════════════════════════

deals = load_deals()

with st.sidebar:
    st.markdown("## 🏢 IC Preparation @ PGIM")
    st.divider()

    # ── Shared Settings ───────────────────────────────────────────────────────
    with st.expander("⚙️  Shared Settings"):
        _ss_path = ROOT / "configs" / "shared_settings.json"
        _ss_data = {}
        if _ss_path.exists():
            try:
                _ss_data = json.loads(_ss_path.read_text(encoding="utf-8"))
            except Exception:
                pass

        _tok_val = st.text_input("Mapbox Token",
                                  value=_ss_data.get("mapbox_token", ""),
                                  type="password",
                                  key="ss_mapbox_token")
        _olb_val = st.text_input("Ollama Base URL",
                                  value=_ss_data.get("ollama_base_url",
                                                      "http://localhost:11434"),
                                  key="ss_ollama_base")

        _omt_val = st.text_input("OneMap API Token",
                                  value=_ss_data.get("onemap_token", ""),
                                  type="password",
                                  key="ss_onemap_token",
                                  help="Optional manual token. Tokens expire ~3 days — prefer the Email/Password below for auto-refresh.")
        _ome_val = st.text_input("OneMap Email",
                                  value=_ss_data.get("onemap_email", ""),
                                  key="ss_onemap_email",
                                  help="OneMap account email (developers.onemap.gov.sg). Used to auto-refresh the token for URA zones + the Location proximity score.")
        _omp_val = st.text_input("OneMap Password",
                                  value=_ss_data.get("onemap_password", ""),
                                  type="password",
                                  key="ss_onemap_password",
                                  help="OneMap account password — stored locally to auto-fetch a fresh token (~3-day expiry).")

        _oai_val = st.text_input("OpenAI API Key",
                                  value=_ss_data.get("openai_api_key", ""),
                                  type="password",
                                  key="ss_openai_key",
                                  help="Required for GPT-4o models. Get yours at platform.openai.com/api-keys")

        _kakao_val = st.text_input("Kakao REST API Key",
                                    value=_ss_data.get("kakao_api_key", ""),
                                    type="password",
                                    key="ss_kakao_key",
                                    help="Required for Korea geocoding. Get yours at developers.kakao.com")

        _gmaps_val = st.text_input("Google Maps API Key",
                                    value=_ss_data.get("google_maps_key", ""),
                                    type="password",
                                    key="ss_google_key",
                                    help="Best for English-romanized Korean (지번) addresses. "
                                         "Enable the Geocoding API at console.cloud.google.com")

        _geo_opts    = ["google", "onemap", "kakao", "mapbox"]
        _geo_labels  = ["🗺️  Google Maps", "🇸🇬  OneMap (Singapore)",
                        "🇰🇷  Kakao (Korea)", "🌍  Mapbox (Global)"]
        _geo_current = _ss_data.get("geocoding_provider", "mapbox").lower()
        _geo_idx     = _geo_opts.index(_geo_current) if _geo_current in _geo_opts else 0
        _geo_sel     = st.radio(
            "Geocoding Provider",
            options=_geo_labels,
            index=_geo_idx,
            key="ss_geo_provider",
            help="Mapbox: global coverage. OneMap: best accuracy for Singapore buildings.",
        )
        _geo_val = _geo_opts[_geo_labels.index(_geo_sel)]

        if st.button("💾  Save Settings", key="ss_save",
                     use_container_width=True):
            try:
                _ss_path.parent.mkdir(parents=True, exist_ok=True)
                _existing = {}
                if _ss_path.exists():
                    try:
                        _existing = json.loads(_ss_path.read_text(encoding="utf-8"))
                    except Exception:
                        pass
                _existing.update({
                    "mapbox_token":       _tok_val.strip(),
                    "ollama_base_url":    _olb_val.strip(),
                    "geocoding_provider": _geo_val,
                    "onemap_token":       _omt_val.strip(),
                    "onemap_email":       _ome_val.strip(),
                    "onemap_password":    _omp_val,
                    "openai_api_key":     _oai_val.strip(),
                    "kakao_api_key":      _kakao_val.strip(),
                    "google_maps_key":    _gmaps_val.strip(),
                })
                _ss_path.write_text(json.dumps(_existing, indent=2), encoding="utf-8")
                st.success("Saved")
            except Exception as e:
                st.error(f"Save failed: {e}")

        # ── Analysis model (also drives online search) ───────────────────────
        st.divider()
        st.markdown(
            "<div style='font-size:11px;letter-spacing:1px;font-weight:700;"
            "opacity:0.7;margin-bottom:6px'>ANALYSIS MODEL</div>",
            unsafe_allow_html=True,
        )
        _ollama_url    = _ollama_base_url()
        _ollama_models = _list_ollama_models(_ollama_url)
        _raw_ollama    = _sort_models_by_speed(_ollama_models) if _ollama_models else []
        _analysis_opts = [_fmt_model(m) for m in _raw_ollama]
        _analysis_opts += [_fmt_model(m) for m in _OPENAI_ANALYSIS_MODELS]
        _analysis_opts += [_NO_LLM_OPTION]
        st.selectbox(
            "🤖 Analysis model",
            _analysis_opts,
            index=0,
            key="sb_analysis_model",
            help="Drives comps scanning, investment rationale, new-deal setup, and "
                 "online search. A GPT selection also runs AI web search on the "
                 "matching OpenAI search model; Ollama / rule-based selections fall "
                 "back to gpt-4o-mini for the online-search step only.",
        )
        _active_analysis = _parse_model_name(
            st.session_state.get("sb_analysis_model",
                                 _analysis_opts[0] if _analysis_opts else ""))
        _is_openai = _active_analysis in _OPENAI_ANALYSIS_MODELS
        if not _ollama_models and not _is_openai:
            st.caption("⚠️ Ollama not detected — start Ollama or pick a GPT model.")
        if _is_openai and not os.environ.get("OPENAI_API_KEY"):
            st.caption("⚠️ OpenAI API key not set — add it in the fields above.")
    st.divider()

    # Top-level nav — two collapsible groups styled like Shared Settings above:
    #   • Analysis Output  → Overview + Comparable Analysis + Investment Rationale
    #   • Deal List        → New Deal + Existing Deals
    # A single ``active_page`` drives the router; the active item is highlighted
    # (primary button) and its group auto-expands.
    _active = st.session_state.setdefault("active_page", "new_deal")

    def _nav_item(label, page_id, key):
        if st.button(label, use_container_width=True, key=key,
                     type=("primary"
                           if st.session_state.get("active_page") == page_id
                           else "secondary")):
            st.session_state["active_page"] = page_id
            st.rerun()

    with st.expander("📊  Analysis Output",
                     expanded=_active in ("overview", "comps", "rationale")):
        _nav_item("📊  Overview",              "overview",  "nav_overview")
        _nav_item("📋  Comparable Analysis",   "comps",     "nav_comps")
        _nav_item("✍️  Investment Rationale",  "rationale", "nav_rationale")

    with st.expander("📁  Deal List",
                     expanded=_active in ("new_deal", "existing")):
        _nav_item("🏗️  New Deal",       "new_deal", "nav_new_deal")
        _nav_item("📁  Existing Deals",  "existing", "nav_existing")
        if st.session_state.get("active_page") == "existing":
            st.markdown("<div style='font-size:11px;letter-spacing:1px;"
                        "font-weight:700;opacity:0.7;margin:6px 0'>"
                        "YOUR DEALS</div>", unsafe_allow_html=True)
            if deals:
                st.selectbox("deals", list(deals.keys()),
                             key="deal_nav", label_visibility="collapsed")
            else:
                st.caption("No deals yet — create one in **New Deal**.")

    # Vision + online-search model selectors have been retired from the sidebar.
    # The single Analysis model in ⚙️ Shared Settings now drives everything,
    # including online search (see _apply_model_overrides). The vision code path
    # (sb_vision_model) remains intact and simply defaults to no override.



# ═════════════════════════════════════════════════════════════════════════════
# MAIN CONTENT  — routes based on sidebar selection
# ═════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────
# ROUTE B:  New Deal  →  Create New form
# ─────────────────────────────────────────────────────────────────────────────

def render_new_deal_form():
    """
    ROUTE B — Two-step wizard for creating a new deal config JSON.

    Step 1 — Collect deal essentials
        The user fills in the minimum required fields: deal name, address,
        asset class, GFA, quality, sale date, and optionally price and cap rate.
        An optional "Upload deal brief" expander lets the user attach a PDF,
        Excel, or text file (or paste raw text).  If a document is provided,
        extract_from_document() uses the LLM to pull out any fields it can find
        before the user types them.

        On clicking "Generate Config Preview", derive_market_fields() sends the
        address and asset class to the LLM, which infers:
          country, country_code, currency, currency_symbol, gfa_unit,
          land_zoning, location descriptor, submarket_keywords, asset_search_keyword,
          broader_market_query
        These derived values are stored in session state and the wizard moves to
        Step 2.

    Step 2 — Review and edit all fields
        An editable data table shows every config field grouped into three
        sections (Property, Deal, Market).  The user can correct any value the
        LLM got wrong.  Submarket keywords are displayed as a comma-separated
        string and split back into a list on save.

        On clicking "Save Config", the dict is assembled into the full deal
        config structure (with llm, mapbox, openai, online_search, parameters
        blocks) and written to configs/deal_config_<DealName>.json.  The wizard
        resets to Step 1 and a success message shows the file path.

    Model usage
        The sidebar model selection drives both extract_from_document() and
        derive_market_fields() calls.  No config file is read for the model —
        the config does not exist yet at this point in the workflow.
    """
    # Model is always taken from the sidebar — no config file dependency
    _active_model = _parse_model_name(st.session_state.get("sb_analysis_model", ""))
    _openai_key   = os.environ.get("OPENAI_API_KEY", "")
    _base_url     = _ollama_base_url()

    if _active_model == _NO_LLM_OPTION:
        llm_cfg = {"provider": "rules"}
    elif _active_model in _OPENAI_ANALYSIS_MODELS:
        llm_cfg = {"provider": "openai", "openai_model": _active_model}
    else:
        # Fall back to first installed Ollama model if sidebar somehow has no selection
        _installed = _list_ollama_models(_base_url)
        _model     = _active_model or (_installed[0] if _installed else "qwen2.5:3b")
        llm_cfg    = {"provider": "ollama",
                      "ollama":   {"base_url": _base_url, "model": _model}}

    st.title("🏗️  New Deal")

    st.caption("Fill in the essentials. "
               "Optionally upload a deal brief — the LLM will extract and pre-fill the fields.")
    if llm_cfg["provider"] == "openai":
        st.warning(
            "**Cloud LLM selected — data-privacy note.** Config generation sends the "
            "**address & asset class** (and any **uploaded deal-brief text**) to the cloud "
            "model (OpenAI) to infer the remaining fields. Don't upload confidential briefs "
            "under a cloud model — switch the sidebar to a **local model (Ollama)** or "
            "**Rule-based (no LLM)** to keep everything on-prem."
        )
    elif llm_cfg["provider"] == "rules":
        st.info(
            "**Rule-based (no LLM) selected** — config fields are derived from the address "
            "+ asset class using built-in rules and Mapbox geocoding only. **Nothing is sent "
            "to any model.** Deal-brief text extraction is disabled in this mode; type the "
            "address and details directly."
        )
    else:
        st.info(
            "Local model selected — config generation runs on-prem (Ollama); "
            "no deal data leaves your machine."
        )
    st.divider()

    # ── Import an existing deal config (skip the form) ───────────────────────
    with st.expander("📥  Import a deal config (JSON) — skip the form"):
        st.caption("Already have a deal config file (exported earlier)? Upload it "
                   "to create the deal directly. Otherwise fill the form below — or "
                   "upload a brief PDF in Step 1 and let the LLM build the config.")
        _imp = st.file_uploader("Upload a deal JSON", type=["json"],
                                key="nd_import_upl")
        if _imp is not None and st.session_state.get("_nd_last_import") != _imp.name:
            import base64 as _b64
            try:
                _data = json.loads(_imp.getvalue().decode("utf-8"))
                if "subject_property" not in _data:
                    raise ValueError("missing 'subject_property'")
                _bundle = _data.pop("_bundle", None)
                _sp = _data["subject_property"]
                _dn = (_sp.get("deal_name") or _sp.get("property_name")
                       or Path(_imp.name).stem)
                import re as _re
                _slug = (_re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", _dn)
                         .strip(" .").replace(" ", "_") or "deal")
                (ROOT / "configs").mkdir(exist_ok=True)
                (ROOT / "configs" / f"deal_config_{_slug}.json").write_text(
                    json.dumps(_data, indent=2), encoding="utf-8")
                if _bundle:   # restore results (records / geo / Excel / map)
                    _out_dir = ROOT / Path(_data.get("output_file",
                                                     "output/x/x.xlsx")).parent
                    _out_dir.mkdir(parents=True, exist_ok=True)
                    for _name, _blob in _bundle.items():
                        try:
                            if "text" in _blob:
                                (_out_dir / _name).write_text(_blob["text"], encoding="utf-8")
                            elif "b64" in _blob:
                                (_out_dir / _name).write_bytes(_b64.b64decode(_blob["b64"]))
                        except Exception:
                            pass
                st.session_state["_nd_last_import"] = _imp.name
                st.session_state["comp_deal"]       = _dn
                st.session_state["deal_nav"]        = _dn
                st.session_state["active_page"]     = "existing"
                st.success(f"✅  Imported **{_dn}**"
                           + (" with results" if _bundle else "")
                           + " — opening it in Existing Deals…")
                st.rerun()
            except Exception as _e:
                st.error(f"Not a valid deal JSON: {_e}")

    # ── Step 1 ─────────────────────────────────────────────────────────────
    st.markdown("##### Step 1 — Deal essentials")

    c1, c2 = st.columns(2)
    with c1:
        qi_deal   = st.text_input("Deal Name *(short)*",
                                   placeholder="e.g. 88 Cecil Street",
                                   key="qi_deal_name")
        qi_addr   = st.text_input("Full Address *(include city & country)*",
                                   placeholder="e.g. 88 Cecil Street, Singapore 069538",
                                   key="qi_address")
    with c2:
        qi_class  = st.selectbox("Asset Class",
                                  ["office","retail","logistics","industrial","mixed-use"],
                                  key="qi_asset_class")
        gfa_c1, gfa_c2 = st.columns([2, 1])
        with gfa_c1:
            qi_gfa = st.text_input("GFA *(number only)*",
                                    placeholder="e.g. 88500",
                                    key="qi_gfa")
        with gfa_c2:
            qi_gfa_unit = st.selectbox("Unit", ["sf", "sqm"],
                                        key="qi_gfa_unit")

    c3, c4, c5, c6, c7 = st.columns(5)
    with c3: qi_quality  = st.text_input("Quality / Grade",
                                          placeholder="Grade A", key="qi_quality")
    with c4: qi_saledt   = st.text_input("Sale Date Label",
                                          placeholder="2025E (Mktg)", key="qi_sale_date")
    with c5: qi_price    = st.text_input("Price (millions, opt.)", key="qi_price")
    with c6: qi_currency = st.selectbox("Currency",
                                         list(_CURRENCY_OPTIONS.keys()),
                                         index=0, key="qi_currency",
                                         help="Currency used to record the price. "
                                              "Auto-detect lets the LLM infer it from the address.")
    with c7: qi_cap      = st.text_input("Cap Rate (decimal, opt.)",
                                          placeholder="0.040", key="qi_cap_rate")

    with st.expander("📄  Optional: Upload deal brief or paste text to pre-fill"):
        d1, d2 = st.columns(2)
        with d1:
            doc_file = st.file_uploader("Upload brief (PDF / Excel / txt)",
                                         type=["pdf","xlsx","txt"], key="nd_doc")
        with d2:
            pasted = st.text_area("Or paste deal info", height=110, key="nd_paste")

    st.write("")
    # Button is enabled if address is filled OR a document/paste is provided
    _has_doc     = doc_file is not None or bool(pasted.strip())
    _has_address = bool(qi_addr.strip())
    if not _has_address and not _has_doc:
        st.caption("ℹ️  Enter an address or upload a deal brief to enable generation.")
    gen_col, _ = st.columns([1, 3])
    with gen_col:
        generate = st.button("🤖  Generate Config Preview",
                             key="nd_generate", type="primary",
                             use_container_width=True,
                             disabled=(not _has_address and not _has_doc))

    if generate:
        with st.spinner(f"Deriving market fields [{_active_model or ollama.get('model','')}] …"):
            try:
                from new_deal import extract_from_document, derive_market_fields

                fields: dict = {
                    "deal_name":               qi_deal.strip() or qi_addr.split(",")[0].strip(),
                    "property_name":           qi_deal.strip(),
                    "address":                 qi_addr.strip(),
                    "asset_class":             qi_class,
                    "gfa_sf":                  qi_gfa.strip(),
                    "quality":                 qi_quality.strip(),
                    "sale_date":               qi_saledt.strip(),
                    "remaining_leasehold_yrs": 0,
                    "price_sgd_m":             qi_price.strip() or None,
                    "price_unit":              "M",
                    "area_unit":               qi_gfa_unit,
                    "ftm_noi_cap_rate":        qi_cap.strip() or None,
                }

                # ── Extract fields from uploaded document / pasted text ────────
                raw_text = ""
                if doc_file is not None:
                    raw_text = _extract_text(doc_file)
                elif pasted.strip():
                    raw_text = pasted.strip()

                if raw_text and llm_cfg["provider"] == "rules":
                    st.info("Rule-based mode: skipping deal-brief text extraction (needs an "
                            "LLM). Using the fields you typed above.")
                elif raw_text:
                    extracted = extract_from_document(raw_text, llm_cfg, _openai_key)
                    for k, v in extracted.items():
                        if v is not None and (fields.get(k) in (None, "", 0)):
                            fields[k] = v

                # ── Address is required by derive_market_fields ───────────────
                if not fields.get("address", "").strip():
                    st.error(
                        "**Address not found.** "
                        "The document did not contain a recognisable address. "
                        "Please type the full address (including city and country) "
                        "in the **Full Address** field above and try again."
                    )
                    st.stop()

                # Auto-fill deal name from address if still blank
                if not fields.get("deal_name", "").strip():
                    fields["deal_name"] = fields["address"].split(",")[0].strip()

                # Get Mapbox token directly from shared_settings.json for geocoding
                _ss_file      = ROOT / "configs" / "shared_settings.json"
                _mapbox_token = ""
                if _ss_file.exists():
                    try:
                        _mapbox_token = json.loads(
                            _ss_file.read_text(encoding="utf-8")).get("mapbox_token", "")
                    except Exception:
                        pass

                derived = derive_market_fields(
                    fields["address"], fields["asset_class"], llm_cfg, _openai_key,
                    mapbox_token=_mapbox_token)
                fields.update(derived)

                if not fields.get("asset_type"):
                    fields["asset_type"] = f"Whole Block ({fields['asset_class'].title()})"

                # Override LLM-derived currency if the user explicitly chose one
                if qi_currency != "Auto-detect from country":
                    code, sym = _CURRENCY_OPTIONS[qi_currency]
                    fields["currency"]        = code
                    fields["currency_symbol"] = sym

                # Always use the user-selected GFA unit (overrides LLM guess)
                fields["gfa_unit"] = qi_gfa_unit

                st.session_state["nd_preview_fields"] = fields
                st.session_state["nd_step"]           = 2
                st.rerun()

            except Exception as e:
                st.error(f"Generation failed: {e}")

    # ── Step 2 — review table ───────────────────────────────────────────────
    if st.session_state.get("nd_step") == 2 and \
       st.session_state.get("nd_preview_fields") is not None:

        st.divider()
        st.markdown("##### Step 2 — Review & edit, then save")
        st.caption("Edit any **Value** cell directly. Submarket Keywords: comma-separated.")

        preview_df = _fields_to_df(st.session_state["nd_preview_fields"])
        edited_df  = st.data_editor(
            preview_df,
            column_config={
                "Section": st.column_config.TextColumn("Section",  width="small",  disabled=True),
                "Field":   st.column_config.TextColumn("Field",    width="medium", disabled=True),
                "_key":    None,
                "Value":   st.column_config.TextColumn("Value",    width="large"),
            },
            hide_index=True,
            use_container_width=True,
            height=700,
            key="nd_editor",
        )

        st.write("")
        sv, rs, _ = st.columns([1, 1, 4])
        with sv:
            save = st.button("💾  Save Config", key="nd_save",
                             type="primary", use_container_width=True)
        with rs:
            if st.button("↩  Start Over", key="nd_reset", use_container_width=True):
                st.session_state["nd_step"]           = 1
                st.session_state["nd_preview_fields"] = None
                st.rerun()

        if save:
            final = _df_to_fields(edited_df)
            missing = [f for f in ("deal_name","address","gfa_sf")
                       if not str(final.get(f,"")).strip()]
            if missing:
                st.error(f"Required: {', '.join(missing)}")
            else:
                try:
                    from new_deal import build_config, save_config
                    config, config_file = build_config(final)
                    save_config(config, str(ROOT / config_file))
                    st.success(f"✅  Saved → `{config_file}`")
                    st.info(f"Select **{final['deal_name']}** in the sidebar "
                            "to view the deal, or go to Comparable Analysis.")
                    st.session_state["nd_step"]           = 1
                    st.session_state["nd_preview_fields"] = None
                    st.rerun()
                except Exception as e:
                    st.error(f"Save failed: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# ROUTE C:  Comparable Analysis
# ─────────────────────────────────────────────────────────────────────────────

def _ensure_adhoc_config(name: str, addr: str, asset_class: str, country: str) -> str:
    """Create/refresh a lightweight ad-hoc subject config so Comparable Analysis can
    run without a full New Deal setup. Named 'adhoc_config_*' so it does NOT appear
    in the Existing Deals list (load_deals only globs 'deal_config*')."""
    import re as _re
    raw  = name or addr.split(",")[0]
    slug = _re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", raw).strip(" .").replace(" ", "_") or "subject"
    path = ROOT / "configs" / f"adhoc_config_{slug}.json"
    cfg  = {}
    if path.exists():
        try:
            cfg = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            cfg = {}
    subj = cfg.get("subject_property", {})
    subj.update({
        "deal_name":     name or slug,
        "property_name": name,
        "address":       addr,
        "asset_class":   asset_class,
        "country_name":  country or "Singapore",
        "country_code":  "sg" if (country or "").strip().lower() == "singapore" else "",
    })
    subj.setdefault("gfa_sf", 0)
    subj.setdefault("gfa_unit", "sf")
    subj.setdefault("quality", "")
    subj.setdefault("remaining_leasehold_yrs", 0)
    cfg["subject_property"] = subj
    cfg.setdefault("output_file",
                   f"output/{slug}/Transaction_Comparables_{slug}.xlsx")
    cfg.setdefault("mapbox", {"style": "streets-v12", "width": 1200, "height": 900,
                              "padding": 100, "pin_size": "l"})
    cfg.setdefault("llm", {"provider": "ollama",
                           "ollama": {"base_url": "http://localhost:11434",
                                      "model": "qwen2.5:3b"}})
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(cfg, indent=2), encoding="utf-8")
    return str(path)


def render_comparable_analysis(deal_name=None, config_path=None):
    """
    ROUTE C — Run comparable analysis for a selected deal.

    Called two ways:
      • Embedded (deal_name + config_path given) — as the 📋 subtab of the
        Analysis Output workspace; the deal is already chosen, so the internal
        subject-source selector is skipped.
      • Standalone (no args) — legacy path with its own subject picker.

    The user picks a deal from the selectbox at the top, then selects one of
    three comp types via a radio button:
      • Asset Sales Comps  — investment transaction comparables
      • Land Sales Comps   — land / site purchase comparables
      • Rent Comps         — leasing / rental comparables

    Each comp type has two tabs:

    📁 Upload Comps
        The user uploads (or reuses a previously configured) input Excel file.
        The backend scan script (scan_input_*.py) is called with the input file:
          1. Ollama auto-detects the column mapping from the header row
          2. Ollama classifies each comp's Location, Quality, and Asset Type
          3. The Bala-adjusted cap rate and price psf are calculated
          4. A formatted 13-column output Excel and Mapbox map PNG are produced
        A temporary config is written with the uploaded file path injected, so
        the original deal config is never modified by the upload.

    🤖 AI Search
        Two sub-options:
          • Internal Database — pick any Excel from Input_files/ and run the
            same scan pipeline (useful for re-running a previously curated file
            on a different deal without re-uploading)
          • Online Search     — calls search_online_*.py, which uses the OpenAI
            web search API (gpt-4o-mini-search-preview) to find real transaction
            data on the internet, then extracts structured fields from the
            articles it finds.  Requires an OpenAI API key.

    After any successful run, _show_results() always renders the latest output
    files (Excel preview table + map) below the buttons, persisting across
    Streamlit re-runs so the results stay visible without re-running.
    """
    if deal_name is None:
        st.title("📋  Comparable Analysis")

    # On the shared cloud link the filesystem is wiped on restart, so remind users
    # to export. Detected via Streamlit Cloud's /mount/src path (no-op on a server).
    if "/mount/" in str(ROOT):
        st.warning("⚠️ **Deals aren't saved on this shared link.** Export your deal "
                   "(📤 below) before you leave to keep it — Import it next time to restore.")

    # Subject can come from an existing deal, OR be keyed in ad-hoc (no full New
    # Deal setup needed) — just a name / location.
    if config_path is not None:
        # Embedded in the Analysis Output workspace — the parent page already
        # picked the deal, so skip the subject-source selector entirely.
        selected = deal_name
        cfg      = load_config(config_path)
        subj     = cfg["subject_property"]
    else:
        _modes = (["📁  Existing deal", "✏️  New subject property"] if deals
                  else ["✏️  New subject property"])
        _mode = st.radio("Subject property source", _modes, horizontal=True,
                         key="comp_subject_mode")

        if _mode.startswith("📁") and deals:
            default_idx = 0
            if "comp_deal" in st.session_state and st.session_state["comp_deal"] in deals:
                default_idx = list(deals.keys()).index(st.session_state["comp_deal"])
            selected    = st.selectbox("Select Deal", list(deals.keys()),
                                       index=default_idx, key="comp_deal_select")
            config_path = deals[selected]
            cfg         = load_config(config_path)
            subj        = cfg["subject_property"]
        else:
            c1, c2 = st.columns(2)
            _ah_name = c1.text_input("Subject property name", key="adhoc_name",
                                     placeholder="e.g. Marina Bay Tower")
            _ah_addr = c2.text_input("Location / address", key="adhoc_addr",
                                     placeholder="e.g. 10 Marina Blvd, Singapore")
            c3, c4 = st.columns(2)
            _ah_class   = c3.selectbox("Asset class",
                                       ["office", "retail", "industrial", "logistics",
                                        "hospitality", "residential", "mixed"],
                                       key="adhoc_class")
            _ah_country = c4.text_input("Country", value="Singapore", key="adhoc_country")
            if _ah_name.strip() or _ah_addr.strip():
                config_path = _ensure_adhoc_config(_ah_name.strip(), _ah_addr.strip(),
                                                   _ah_class, _ah_country.strip())
                cfg  = load_config(config_path)
                subj = cfg["subject_property"]
            else:
                # No subject typed yet — leave config unset but still fall through
                # so the Import panel below renders (a first-time visitor with zero
                # deals must be able to restore a saved deal from a JSON file).
                config_path = None
                cfg = subj = None

    # ── Export deal (portable per-user storage; no server DB needed) ──────────
    # On the cloud the filesystem is wiped on restart, so users keep their own deal
    # as a JSON file on their machine. Export downloads it here; Import lives on the
    # New Deal page (upload a config JSON, or a brief PDF for the LLM to extract).
    with st.expander("📤 Export deal (save as a file)"):
        import base64 as _b64
        _scope = st.radio(
            "Export scope",
            ["Subject only", "Subject + results (full restore)"],
            key="export_scope",
            help="'Subject only' saves just the deal setup (re-run comps after "
                 "import). 'Subject + results' also bundles the generated comps, "
                 "Excel and map so importing restores the finished analysis.")
        try:
            if config_path is None:
                raise ValueError("no subject yet")
            _cfg_data = load_config(config_path)
            _full = _scope.startswith("Subject +")
            if _full:
                _out_dir = ROOT / Path(_cfg_data.get("output_file",
                                                     "output/x/x.xlsx")).parent
                _bundle = {}
                if _out_dir.exists():
                    for _f in _out_dir.iterdir():
                        if not _f.is_file():
                            continue
                        if _f.suffix.lower() == ".json":
                            _bundle[_f.name] = {"text": _f.read_text(encoding="utf-8")}
                        elif _f.suffix.lower() in (".xlsx", ".png"):
                            _bundle[_f.name] = {"b64": _b64.b64encode(_f.read_bytes()).decode()}
                _cfg_data["_bundle"] = _bundle
            _out_name = Path(config_path).stem + ("_full" if _full else "") + ".json"
            st.download_button("⬇️  Export", json.dumps(_cfg_data, indent=2),
                               file_name=_out_name, mime="application/json",
                               use_container_width=True)
        except Exception:
            st.caption("Nothing to export yet — pick a deal above, or import one "
                       "on the **New Deal** page.")

    # Subject gate — must come AFTER the Import panel so importing is always reachable.
    if subj is None:
        st.info("Enter a subject property **name or location** to begin, "
                "or import a saved deal above.")
        return

    # Inline subject info strip
    i1, i2, i3, i4 = st.columns(4)
    i1.markdown(f"**📍** {subj.get('address') or '—'}")
    i2.markdown(f"**Class:** {(subj.get('asset_class') or '—').title()}")
    i3.markdown(f"**GFA:** {int(subj.get('gfa_sf', 0) or 0):,} {subj.get('gfa_unit', 'sf').upper()}")
    i4.markdown(f"**Quality:** {subj.get('quality') or '—'}")

    st.divider()

    comp_type = st.radio(
        "Comparable type",
        ["🏢  Asset Sales Comps", "🌍  Land Sales Comps", "📄  Rent Comps"],
        horizontal=True, key="comp_type_radio",
    )
    is_asset_sales = comp_type.startswith("🏢")
    is_land_sales  = comp_type.startswith("🌍")
    is_rent        = comp_type.startswith("📄")
    st.write("")

    sub_upload, sub_ai = st.tabs(["📁  Upload Comps", "🤖  AI Search"])

    # ── Upload Comps ──────────────────────────────────────────────────────────
    with sub_upload:
        st.subheader("Upload Comparable Documents")

        if is_asset_sales:
            configured_key       = "input_file"
            configured_pdf_key   = "input_pdf_file"
            configured_image_key = "input_image_file"
            upload_script        = "scan_input_sales_comps.py"
            upload_flags         = ["--map"]
            result_prefix        = "Transaction_Comparables"
            st.caption("Asset sales comps — upload Excel, PDF, and/or image (records are merged).")
        elif is_land_sales:
            configured_key       = "land_input_file"
            configured_pdf_key   = "land_input_pdf_file"
            configured_image_key = "land_input_image_file"
            upload_script        = "scan_input_land_comps.py"
            upload_flags         = ["--map"]
            result_prefix        = "Land_Sale_Comps"
            st.caption("Land sale comps — upload Excel, PDF, and/or image (records are merged).")
        else:  # rent
            configured_key       = "rent_input_file"
            configured_pdf_key   = "rent_input_pdf_file"
            configured_image_key = "rent_input_image_file"
            upload_script        = "scan_input_rent_comps.py"
            upload_flags         = ["--map"]
            result_prefix        = "Rent_Comps"
            st.caption("Rent comps — upload Excel, PDF, and/or image (records are merged).")

        configured_file       = cfg.get(configured_key)
        configured_pdf_file   = cfg.get(configured_pdf_key)
        configured_image_file = cfg.get(configured_image_key)

        # ── Input files (single uploader: Excel, PDF, and/or Image) ─────────
        with st.container(border=True):
            st.markdown("**Input Files** — upload Excel, PDF, and/or image screenshot")
            uploaded_files = st.file_uploader(
                "Upload Excel (.xlsx), PDF (.pdf), and/or image (.png / .jpg)",
                type=["xlsx", "pdf", "png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key=f"up_{comp_type}",
            )
            clear_excel = False
            clear_pdf   = False
            clear_image = False
            if configured_file:
                col_info, col_clr = st.columns([6, 1])
                if isinstance(configured_file, list):
                    _xl_label = ", ".join(f"`{p}`" for p in configured_file)
                    col_info.info(f"**Excels configured ({len(configured_file)}):** {_xl_label}")
                else:
                    col_info.info(f"**Excel configured:** `{configured_file}`")
                clear_excel = col_clr.checkbox("Remove", key=f"clr_xl_{comp_type}",
                                               help="Remove this Excel from config")
            if configured_pdf_file:
                col_info, col_clr = st.columns([6, 1])
                if isinstance(configured_pdf_file, list):
                    _pdf_label = ", ".join(f"`{p}`" for p in configured_pdf_file)
                    col_info.info(f"**PDFs configured ({len(configured_pdf_file)}):** {_pdf_label}")
                else:
                    col_info.info(f"**PDF configured:** `{configured_pdf_file}`")
                clear_pdf = col_clr.checkbox("Remove", key=f"clr_pdf_{comp_type}",
                                             help="Remove this PDF from config")
            if configured_image_file:
                col_info, col_clr = st.columns([6, 1])
                if isinstance(configured_image_file, list):
                    _img_label = ", ".join(f"`{p}`" for p in configured_image_file)
                    col_info.info(f"**Images configured ({len(configured_image_file)}):** {_img_label}")
                else:
                    col_info.info(f"**Image configured:** `{configured_image_file}`")
                clear_image = col_clr.checkbox("Remove", key=f"clr_img_{comp_type}",
                                               help="Remove this image from config")

        # Split uploaded files by type
        _img_exts = {".png", ".jpg", ".jpeg"}
        uploaded_excels = [f for f in uploaded_files if f.name.lower().endswith(".xlsx")]
        uploaded_excel  = uploaded_excels[0] if uploaded_excels else None  # backwards-compat alias
        uploaded_pdfs   = [f for f in uploaded_files if f.name.lower().endswith(".pdf")]
        uploaded_pdf    = uploaded_pdfs[0] if uploaded_pdfs else None      # backwards-compat alias
        uploaded_images = [f for f in uploaded_files if Path(f.name).suffix.lower() in _img_exts]
        uploaded_image  = uploaded_images[0] if uploaded_images else None  # backwards-compat alias

        # ── Manual input table ────────────────────────────────────────────────
        # Use when extraction fails, GFA is missing, or records need correction.
        # Saves as a temp Excel that feeds through the normal Excel pipeline.
        import pandas as _pd
        with st.expander("✏️  Enter or paste records manually"):
            if is_asset_sales:
                _manual_cols = ["Property", "Address", "Sale Date",
                                "Land Zoning", "Remaining Leasehold (Y)",
                                "GFA (SF)", "Price (SGD M)", "NPI Yield (%)", "Sale Type"]
            elif is_land_sales:
                _manual_cols = ["Site Name", "Address", "Date of Launch",
                                "Land Zoning", "Tenure",
                                "Site Area (SF)", "Max GFA (SF)",
                                "Price (SGD M)", "Price psf ppr", "Sale Type", "Comment"]
            else:  # rent
                _manual_cols = ["Property", "Address", "Lease Date",
                                "NLA (SF)", "Asking Rent", "Eff. Rent",
                                "Lease Term (Yrs)", "Rent-Free (Mths)", "Asset Type"]

            _mk = f"manual_df_{comp_type}"
            if _mk not in st.session_state:
                st.session_state[_mk] = _pd.DataFrame(columns=_manual_cols)

            manual_df = st.data_editor(
                st.session_state[_mk],
                num_rows="dynamic",
                use_container_width=True,
                key=f"manual_editor_{comp_type}",
            )
            _manual_rows = manual_df.dropna(how="all")
            if len(_manual_rows):
                st.caption(f"✓ {len(_manual_rows)} row(s) — will be merged with any uploaded files above.")
            else:
                st.caption("Add rows using the ＋ at the bottom of the table.")

        _has_manual = len(manual_df.dropna(how="all")) > 0

        # Remaining tenure override — shown for asset sales and land comps only
        remaining_yrs_ui = None
        if not is_rent:
            subj_yrs_cfg = int(cfg.get("subject_property", {})
                                .get("remaining_leasehold_yrs", 0))
            col_t, _ = st.columns([2, 3])
            with col_t:
                remaining_yrs_ui = st.number_input(
                    "Subject Remaining Tenure (yrs) — 0 = Freehold",
                    min_value=0, max_value=999,
                    value=subj_yrs_cfg,
                    step=1,
                    key=f"rem_yrs_{comp_type}",
                    help="Used for Bala Table adjustment in Adj. Cap Rate / Adj. Price."
                )

        active_input = configured_file or configured_pdf_file or configured_image_file
        tmp_cfg_path = None
        patched      = dict(cfg)
        needs_patch  = False

        # Apply "Remove" checkboxes — blank out the key in the patched config
        # and also update active_input so the Run button disables correctly.
        if clear_excel:
            patched[configured_key] = None
            if active_input == configured_file:
                active_input = configured_pdf_file or configured_image_file
            needs_patch = True
            st.warning(f"Excel will be removed from config on next run.")
        if clear_pdf:
            patched[configured_pdf_key] = None
            if active_input == configured_pdf_file:
                active_input = configured_file or configured_image_file
            needs_patch = True
            st.warning(f"PDF will be removed from config on next run.")
        if clear_image:
            patched[configured_image_key] = None
            if active_input == configured_image_file:
                active_input = configured_file or configured_pdf_file
            needs_patch = True
            st.warning(f"Image will be removed from config on next run.")

        # New uploads APPEND to any already-configured inputs of the same kind
        # (deduped by path), rather than replacing them. The merge base is
        # ``patched`` so a "Remove" ticked above clears first, then the new
        # upload appends to the emptied list.
        def _append_saved(key, saved_rels):
            _merged = _input_list(patched.get(key))
            for _rel in saved_rels:
                if _rel not in _merged:
                    _merged.append(_rel)
            # Single item → plain string (backwards compat); multiple → list
            patched[key] = _merged[0] if len(_merged) == 1 else _merged
            return _merged

        if uploaded_excels:
            _saved_excels = []
            for _uf in uploaded_excels:
                _sp = ROOT / "Input_files" / _uf.name
                _sp.parent.mkdir(parents=True, exist_ok=True)
                _sp.write_bytes(_uf.getvalue())
                _saved_excels.append(str(_sp.relative_to(ROOT)))
                st.success(f"Excel saved → `{_saved_excels[-1]}`")
            _merged_xl = _append_saved(configured_key, _saved_excels)
            active_input = _merged_xl[0]
            needs_patch = True

        if uploaded_pdfs:
            _saved_pdfs = []
            for _uf in uploaded_pdfs:
                _pp = ROOT / "Input_files" / _uf.name
                _pp.parent.mkdir(parents=True, exist_ok=True)
                _pp.write_bytes(_uf.getvalue())
                _saved_pdfs.append(str(_pp.relative_to(ROOT)))
                st.success(f"PDF saved → `{_saved_pdfs[-1]}`")
            _merged_pdf = _append_saved(configured_pdf_key, _saved_pdfs)
            active_input = active_input or _merged_pdf[0]
            needs_patch = True

        if uploaded_images:
            _saved_images = []
            for _uf in uploaded_images:
                _ip = ROOT / "Input_files" / _uf.name
                _ip.parent.mkdir(parents=True, exist_ok=True)
                _ip.write_bytes(_uf.getvalue())
                _saved_images.append(str(_ip.relative_to(ROOT)))
                st.success(f"Image saved → `{_saved_images[-1]}`")
            _merged_img = _append_saved(configured_image_key, _saved_images)
            active_input = active_input or _merged_img[0]
            needs_patch = True

        if _has_manual:
            _raw_slug = cfg.get("subject_property", {}).get("deal_name",
                         cfg.get("subject_property", {}).get("property_name", "deal"))
            # Strip characters Windows forbids in file paths (< > : " / \ | ? *) —
            # an unsanitised name causes [Errno 22] Invalid argument on Windows.
            _deal_slug = "".join("_" if _c in '<>:"/\\|?*' else _c
                                 for _c in _raw_slug).strip(" .").replace(" ", "_") or "deal"
            _manual_path = ROOT / "Input_files" / f"_manual_{comp_type}_{_deal_slug}.xlsx"
            _manual_path.parent.mkdir(parents=True, exist_ok=True)
            manual_df.dropna(how="all").to_excel(str(_manual_path), index=False)
            _manual_rel = str(_manual_path.relative_to(ROOT))
            st.success(f"Manual data saved → `{_manual_rel}`  ({len(manual_df.dropna(how='all'))} rows)")
            if not uploaded_excel:
                patched[configured_key] = _manual_rel
            active_input = active_input or _manual_rel
            needs_patch = True

        if remaining_yrs_ui is not None:
            patched["subject_property"] = {
                **patched.get("subject_property", {}),
                "remaining_leasehold_yrs": int(remaining_yrs_ui),
            }
            needs_patch = True

        if needs_patch:
            tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".json",
                                              dir=ROOT/"configs", delete=False,
                                              encoding="utf-8")
            json.dump(patched, tmp); tmp.close()
            tmp_cfg_path = tmp.name

            # Persist input-file changes to the REAL config so uploads ACCUMULATE
            # (append) across runs/sessions, and "Remove" ticks stick. Only the
            # three input keys are copied over from ``patched`` — other runtime
            # patches (e.g. tenure override) stay in the temp config only.
            _uploaded_any = bool(uploaded_excels or uploaded_pdfs or uploaded_images)
            if _uploaded_any or clear_excel or clear_pdf or clear_image:
                _save = load_config(config_path)
                for _k in (configured_key, configured_pdf_key, configured_image_key):
                    _val = patched.get(_k)
                    if _val:
                        _save[_k] = _val
                    else:
                        _save.pop(_k, None)
                Path(config_path).write_text(
                    json.dumps(_save, indent=2, ensure_ascii=False), encoding="utf-8")
                # Rerun only on Remove, to refresh the "configured" display. Do
                # NOT rerun on upload — the uploader still holds the files, which
                # would re-trigger this block and loop.
                if clear_excel or clear_pdf or clear_image:
                    st.rerun()

        st.write("")
        _name_only = st.checkbox(
            "Property name only — keep a comp if it has a name, even without price/area",
            value=False, key=f"name_only_{comp_type}",
            help="When ticked, a row qualifies on its property name alone — price, area "
                 "and other field requirements are skipped. Useful for a curated list of "
                 "names to geocode/map. Applies to Excel, PDF and image inputs.")
        _run_flags = upload_flags + (["--name-only"] if _name_only else [])
        if st.button("▶  Run Analysis", key="run_upload", type="primary",
                     disabled=(active_input is None)):
            use = tmp_cfg_path or config_path
            if _run_script(upload_script, use, _run_flags):
                st.success("✅  Done!")
            if tmp_cfg_path:
                try: os.unlink(tmp_cfg_path)
                except: pass

        _upload_log_key = f"_runlog_{upload_script}"
        if st.session_state.get(_upload_log_key):
            with st.expander("📋 Last run log", expanded=False):
                if st.button("🗑️ Clear log", key="clear_upload_log"):
                    st.session_state.pop(_upload_log_key, None)
                    st.session_state.pop(f"{_upload_log_key}_err", None)
                    st.rerun()
                st.code(st.session_state[_upload_log_key], language="text")
                if st.session_state.get(f"{_upload_log_key}_err"):
                    st.code(st.session_state[f"{_upload_log_key}_err"], language="text")

        # Always render latest output files — persists across Streamlit re-runs
        _show_results(config_path, result_prefix, context="upload",
                      script=upload_script)

    # ── AI Search ─────────────────────────────────────────────────────────────
    with sub_ai:
        st.subheader("AI-Powered Comparable Search")

        st.caption("Find comparables from your internal files or the web.")
        source = st.radio("Source", ["🗄️  Internal Database", "🌐  Online Search"],
                          horizontal=True, key="ai_source")
        online = source.startswith("🌐")

        st.write("")

        if not online:
            # ── Folder to browse — your comp database (searched recursively) ──────
            # Defaults to the last-used folder (persisted in shared_settings), else
            # Input_files. Point it at your network comp DB, e.g.
            #   K:\Transactions\_Comp Database\IC Comps Inbox
            _ssp = ROOT / "configs" / "shared_settings.json"
            def _ss_read():
                try:    return json.loads(_ssp.read_text(encoding="utf-8"))
                except Exception: return {}
            _default_db = _ss_read().get("comp_db_dir", "") or str(ROOT / "Input_files")
            db_dir_str = st.text_input(
                "📁  Comp database folder",
                value=st.session_state.get("db_dir", _default_db), key="db_dir",
                help=r"Folder searched recursively for Excel/PDF comps. "
                     r"e.g. K:\Transactions\_Comp Database\IC Comps Inbox  "
                     "(local only — the cloud can't see network drives).")
            db_base = Path(db_dir_str.strip().strip('"'))

            _c1, _c2 = st.columns([4, 1])
            _query = _c1.text_input("🔎  Search files", key="db_search",
                                    placeholder="part of a file name or subfolder…",
                                    label_visibility="collapsed")

            @st.cache_data(show_spinner="Scanning folder…")
            def _list_comp_files(base_str: str):
                base = Path(base_str)
                if not base.exists():
                    return None
                exts = {".xlsx", ".xls", ".pdf"}
                out = []
                try:
                    for p in base.rglob("*"):
                        if (p.is_file() and p.suffix.lower() in exts
                                and not p.name.startswith(("~$", "."))):
                            out.append(p)
                        if len(out) >= 8000:
                            break
                except Exception:
                    return None
                return sorted(out, key=lambda p: str(p).lower())

            if _c2.button("🔄  Rescan", key="db_rescan"):
                _list_comp_files.clear()

            all_files = _list_comp_files(str(db_base))
            chosen_path = None
            if all_files is None:
                st.warning(f"Folder not found or unreadable: `{db_base}`  "
                           "(on the cloud, use `Input_files`).")
            elif not all_files:
                st.info("No Excel/PDF files found in that folder.")
            else:
                q = _query.strip().lower()
                shown = [p for p in all_files if q in str(p).lower()] if q else all_files
                def _label(p):
                    try:    return str(p.relative_to(db_base))
                    except ValueError: return p.name
                opts = {_label(p): p for p in shown[:1000]}
                st.caption(f"{len(shown)} file(s)"
                           + (f" matching '{_query}'" if q else "")
                           + (" — showing first 1000" if len(shown) > 1000 else ""))
                if opts:
                    chosen_name = st.selectbox("Reference file", list(opts.keys()),
                                               key="db_file")
                    chosen_path = str(opts[chosen_name])
                else:
                    st.info("No files match your search.")

            # Determine script/prefix + config keys (Excel vs PDF) before the button.
            if is_asset_sales:
                _xl_key, _pdf_key, db_script, db_flags, db_prefix = (
                    "input_file", "input_pdf_file", "scan_input_sales_comps.py",
                    [], "Transaction_Comparables")
            elif is_land_sales:
                _xl_key, _pdf_key, db_script, db_flags, db_prefix = (
                    "land_input_file", "land_input_pdf_file", "scan_input_land_comps.py",
                    ["--map"], "Land_Sale_Comps")
            else:
                _xl_key, _pdf_key, db_script, db_flags, db_prefix = (
                    "rent_input_file", "rent_input_pdf_file", "scan_input_rent_comps.py",
                    ["--map"], "Rent_Comps")

            if st.button("▶  Search Database", key="run_db", type="primary",
                         disabled=(chosen_path is None)):
                # Remember the folder for next time (best-effort; on-prem only).
                try:
                    _ss = _ss_read(); _ss["comp_db_dir"] = str(db_base)
                    _ssp.write_text(json.dumps(_ss, indent=2), encoding="utf-8")
                except Exception:
                    pass
                _is_pdf = chosen_path.lower().endswith(".pdf")
                patched = {**cfg, (_pdf_key if _is_pdf else _xl_key): chosen_path}
                tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".json",
                                                  dir=ROOT/"configs", delete=False,
                                                  encoding="utf-8")
                json.dump(patched, tmp); tmp.close()
                ok = _run_script(db_script, tmp.name, db_flags)
                try: os.unlink(tmp.name)
                except: pass
                if ok:
                    st.success("✅  Done!")

            # Always render latest output — persists across re-runs
            _show_results(config_path, db_prefix, context="db",
                          script=db_script)

        else:
            # Online search (Asset Sales or Rent only)
            api_key = (os.environ.get("OPENAI_API_KEY")
                       or cfg.get("openai", {}).get("api_key") or "")
            if not api_key:
                st.warning("OpenAI API key required.")
                k = st.text_input("OpenAI API key", type="password",
                                  placeholder="sk-...", key="oai_key")
                if k:
                    os.environ["OPENAI_API_KEY"] = k
                    api_key = k
            else:
                st.success("✅  OpenAI API key detected.")

            refresh = st.checkbox("Force refresh (ignore cache)", key="force_refresh")

            # Determine script/prefix (and the search-config key to patch sources into).
            if is_asset_sales:
                ol_script, ol_prefix, _sc_key = ("search_online_sales_comps.py",
                                                 "Online_Comparables", "online_search")
                _comp_type = "sales"
            elif is_land_sales:
                ol_script, ol_prefix, _sc_key = ("search_online_land_comps.py",
                                                 "Online_Land_Comps", "land_search")
                _comp_type = "land"
            else:
                ol_script, ol_prefix, _sc_key = ("search_online_rent_comps.py",
                                                 "Online_Rent_Comps", "rent_search")
                _comp_type = "rent"
            ol_flags = ["--map"] + (["--refresh"] if refresh else [])

            # ── Data sources (grounded official sources + cited web search) ──────
            _market = (cfg.get("country_code", "") or "sg").lower()
            try:
                from sources.registry import available as _sources_available
                _src_opts = _sources_available(_market, _comp_type)
            except Exception:
                _src_opts = ["web_search"]
            _src_labels = {
                "web_search":      "🌐 Web search (GPT, cited)",
                "ura_pmi":         "🏛️ URA PMI — commercial transactions (SG, free)",
                "ura_pmi_rental":  "🏛️ URA PMI — commercial rents by street (SG, free)",
                "ura_gls":         "🏛️ URA GLS — land tenders (SG, free)",
                "broker_reports":  "📑 Broker reports — Savills / C&W / CBRE (SG)",
            }
            _selected_sources = st.multiselect(
                "Data sources", _src_opts, default=_src_opts,
                format_func=lambda s: _src_labels.get(s, s), key="online_sources",
                help="Grounded official sources (e.g. URA) are combined with the "
                     "citation-backed web search. Deselect to narrow.")
            if not _selected_sources:
                _selected_sources = ["web_search"]

            if st.button("▶  Search Online", key="run_online", type="primary",
                         disabled=not api_key):
                # Patch the chosen sources into a temp config (cache auto-busts on change).
                _patched = {**cfg}
                _sc = {**_patched.get(_sc_key, {})}
                _sc["sources"] = _selected_sources
                _patched[_sc_key] = _sc
                _tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".json",
                                                   dir=ROOT / "configs", delete=False,
                                                   encoding="utf-8")
                json.dump(_patched, _tmp); _tmp.close()
                try:
                    if _run_script(ol_script, _tmp.name, ol_flags):
                        st.success("✅  Online search complete!")
                finally:
                    try:
                        os.unlink(_tmp.name)
                    except OSError:
                        pass

            # Always render latest output — persists across re-runs
            _show_results(config_path, ol_prefix, context="online",
                          script=ol_script)


# ─────────────────────────────────────────────────────────────────────────────
# ROUTE D:  Investment Rationale
# ─────────────────────────────────────────────────────────────────────────────

_SHOW_RATIONALE_TIMING = True   # set False to hide "Generated on … took …s" caption

def render_investment_rationale(deal_name=None, config_path=None):
    """
    ROUTE D — Generate and refine the 3-section investment rationale.

    Called embedded (deal_name + config_path given) as the ✍️ subtab of the
    Analysis Output workspace, or standalone (no args) with its own deal picker.

    Overview
    --------
    The user selects a deal, ticks which PDF market reports to include, adds
    optional analyst notes, and clicks Generate.  The backend runs a two-stage
    LLM pipeline (generate_investment_rationale.py):

      Stage 1 — Extract (cached)
          Each selected PDF is read and the LLM extracts structured market
          intelligence (vacancy, rental trends, cap values, demand drivers, etc.).
          Results are cached by file hash in Input_files/market_reports/cache/.
          Unchanged reports are served instantly on subsequent runs — no LLM call.
          The checkbox "♻️ Re-extract reports" forces a fresh extraction.

      Stage 2 — Generate (two separate LLM calls)
          Call 1: writes the three prose sections using anonymised source labels
                  ("Research Report 1", "Research Report 2", …) so that no PDF
                  filename can appear in the body text.
          Call 2: audits every claim against the real source data and produces
                  a citation JSON with source filename, page reference, and
                  supporting text.  This goes into Source_Audit.xlsx only.

    Output display
    --------------
    The finished markdown rationale is displayed inline and can be downloaded
    as a .md file or converted to a Word .docx via _markdown_to_docx().
    The Source_Audit.xlsx is available as a download and also previewed in a
    collapsible expander below the rationale.

    Refine panel (🔄 Refine This Output)
    ----------------------------------------
    After generation, the supervisor can expand the Refine panel, type specific
    feedback (e.g. "Section 1 title too generic, focus on e-commerce angle"),
    and click "Regenerate with Changes".  Only Stage 2 (writing) re-runs —
    extraction stays cached — so local model re-runs take ~2-4 min and GPT
    re-runs take <30 s.  The refinement feedback is written to a temp file and
    passed to the backend via --refinement-file flag, where it is injected into
    the prompt as a high-priority REFINEMENT INSTRUCTIONS block.
    """
    if deal_name is None:
        st.title("✍️  Investment Rationale")
        st.caption(
            "Select a deal, tick the market reports you want to use, "
            "add any analyst notes, then generate a 3-section investment rationale."
        )
    st.warning(
        "**LLM-generated — be mindful of data privacy.** This feature sends the selected "
        "market-report text and deal details to the chosen analysis model to extract "
        "insights and write the memo. If a **cloud model (e.g. GPT)** is selected, that "
        "content leaves your machine — only use reports cleared for external processing. "
        "Select a **local model (Ollama)** to keep everything on-prem. Always verify the "
        "output against the source PDFs."
    )
    st.divider()

    if not deals:
        st.info("No deals yet — create one in **New Deal** first.")
        return

    # ── Deal selector ─────────────────────────────────────────────────────────
    if config_path is not None:
        # Embedded in Analysis Output — deal already chosen by the parent page.
        selected    = deal_name
        cfg         = load_config(config_path)
        subj        = cfg["subject_property"]
    else:
        selected = st.selectbox(
            "Select Deal", list(deals.keys()),
            key="ir_deal_select",
        )
        config_path = deals[selected]
        cfg         = load_config(config_path)
        subj        = cfg["subject_property"]

    # Inline deal strip
    i1, i2, i3, i4 = st.columns(4)
    i1.markdown(f"**📍** {subj.get('address', '—')}")
    i2.markdown(f"**Class:** {subj.get('asset_class', '—').title()}")
    i3.markdown(f"**GFA:** {int(subj.get('gfa_sf', 0)):,} {subj.get('gfa_unit','sf').upper()}")
    i4.markdown(f"**Quality:** {subj.get('quality', '—')}")
    st.write("")

    # ── Market reports picker ─────────────────────────────────────────────────
    reports_dir  = ROOT / "Input_files" / "market_reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    available    = sorted(p for p in reports_dir.glob("*.pdf") if not p.name.startswith("."))

    st.markdown("#### 📂  Market Reports")

    # Upload new reports
    with st.expander("⬆️  Upload new report(s)", expanded=False):
        uploaded_rpts = st.file_uploader(
            "Upload PDF market reports",
            type=["pdf"],
            accept_multiple_files=True,
            key="ir_upload_reports",
        )
        if uploaded_rpts:
            for f in uploaded_rpts:
                dest = reports_dir / f.name
                dest.write_bytes(f.getvalue())
                st.success(f"Saved → `Input_files/market_reports/{f.name}`")
            # Refresh the available list after upload
            available = sorted(p for p in reports_dir.glob("*.pdf")
                               if not p.name.startswith("."))

    if not available:
        st.warning(
            "No PDF reports in `Input_files/market_reports/`.  "
            "Upload at least one report above."
        )
        return

    st.caption(f"{len(available)} report(s) available — tick those to include:")
    selected_reports = []
    for rpt in available:
        # Check if cached (gives user visibility)
        cache_dir = ROOT / "Input_files" / "market_reports" / "cache"
        cached    = list(cache_dir.glob(f"{rpt.stem}_*.json")) if cache_dir.exists() else []
        label     = f"{rpt.name}  ✅ *(extracted)*" if cached else rpt.name
        if st.checkbox(label, value=True, key=f"ir_rpt_{rpt.name}"):
            selected_reports.append(str(rpt))

    st.write("")

    # ── Analyst notes ──────────────────────────────────────────────────────────
    st.markdown("#### 📝  Analyst Notes *(optional)*")
    analyst_notes = st.text_area(
        "Add any additional context, deal highlights, or specific angles to emphasise",
        height=130,
        placeholder=(
            "e.g. The vendor is motivated; deal expected to close at a 20bps discount to market.\n"
            "Highlight the long WALE and recent anchor tenant renewal."
        ),
        key="ir_notes",
    )

    # ── Options ───────────────────────────────────────────────────────────────
    col_opt1, col_opt2, _ = st.columns([1, 1, 3])
    with col_opt1:
        force_refresh = st.checkbox("♻️  Re-extract reports (ignore cache)", key="ir_refresh")

    st.write("")

    # ── Generate button ───────────────────────────────────────────────────────
    gen_disabled = len(selected_reports) == 0
    if st.button(
        "🤖  Generate Investment Rationale",
        type="primary",
        disabled=gen_disabled,
        use_container_width=False,
        key="ir_generate",
    ):
        # Write analyst notes to a temp file if provided
        notes_file_path = ""
        if analyst_notes.strip():
            tf = tempfile.NamedTemporaryFile(
                mode="w", suffix=".txt", delete=False,
                dir=ROOT / "configs", encoding="utf-8",
            )
            tf.write(analyst_notes)
            tf.close()
            notes_file_path = tf.name

        flags = []
        if selected_reports:
            flags += ["--reports"] + selected_reports
        if force_refresh:
            flags.append("--refresh")
        if notes_file_path:
            flags += ["--notes-file", notes_file_path]

        _t0 = time.time()
        ok  = _run_script("generate_investment_rationale.py", config_path, flags)
        _elapsed = time.time() - _t0

        # Clean up temp notes file
        if notes_file_path:
            try:
                os.unlink(notes_file_path)
            except Exception:
                pass

        if ok:
            if _SHOW_RATIONALE_TIMING:
                _meta_out = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
                try:
                    import json as _jt, datetime as _dt
                    (_meta_out / "Investment_Rationale_meta.json").write_text(
                        _jt.dumps({
                            "generated_at": _dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "elapsed_s":    round(_elapsed, 1),
                        }), encoding="utf-8")
                except Exception:
                    pass
            st.success("✅  Rationale generated!")
            st.rerun()

    # ── Display existing / freshly generated rationale ────────────────────────
    out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
    rationale_file = out_dir / "Investment_Rationale.md"

    if rationale_file.exists():
        st.divider()
        st.markdown("#### 📄  Investment Rationale")

        if _SHOW_RATIONALE_TIMING:
            _meta_path = out_dir / "Investment_Rationale_meta.json"
            if _meta_path.exists():
                try:
                    import json as _jt
                    _meta = _jt.loads(_meta_path.read_text(encoding="utf-8"))
                    _gen  = _meta.get("generated_at", "")
                    _secs = _meta.get("elapsed_s")
                    _cap  = f"Generated {_gen}"
                    if _secs is not None:
                        _cap += f"  ·  took {_secs}s"
                    st.caption(_cap)
                except Exception:
                    pass

        rationale_text = rationale_file.read_text(encoding="utf-8")
        # Escape '$' so Streamlit doesn't render "S$353 … S$1.83" as LaTeX math
        # (it treats text between two '$' as a formula). The downloaded .md keeps
        # the raw '$'.
        st.markdown(rationale_text.replace("$", "\\$"))

        st.write("")
        dl1, dl2, dl3, _ = st.columns([1, 1, 1, 2])
        with dl1:
            st.download_button(
                "⬇️  Download (.md)",
                data=rationale_text.encode("utf-8"),
                file_name=f"Investment_Rationale_{selected.replace(' ','_')}.md",
                mime="text/markdown",
                key="ir_dl_md",
                use_container_width=True,
            )
        with dl2:
            try:
                docx_bytes = _markdown_to_docx(rationale_text)
                st.download_button(
                    "⬇️  Download (.docx)",
                    data=docx_bytes,
                    file_name=f"Investment_Rationale_{selected.replace(' ','_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="ir_dl_docx",
                    use_container_width=True,
                )
            except ImportError:
                st.caption("Install `python-docx` for Word download.")

        # Source audit Excel
        audit_file = out_dir / "Source_Audit.xlsx"
        if audit_file.exists():
            with open(audit_file, "rb") as fh:
                audit_bytes = fh.read()
            with dl3:
                st.download_button(
                    "⬇️  Source Audit (.xlsx)",
                    data=audit_bytes,
                    file_name=f"Source_Audit_{selected.replace(' ','_')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="ir_dl_audit",
                    use_container_width=True,
                )

            # ── Refine panel ──────────────────────────────────────────────────
            st.write("")
            with st.expander("🔄  Refine This Output", expanded=False):
                st.caption(
                    "PDF extraction is cached — only the writing step re-runs. "
                    "Typically **2–4 min** for local models, **<30 s** for GPT."
                )
                refinement_feedback = st.text_area(
                    "What to change in the next draft",
                    height=120,
                    placeholder=(
                        "e.g. Section 1 title is too generic — focus on the e-commerce "
                        "demand angle specifically.\n"
                        "Section 2 should mention the proximity to the port more.\n"
                        "Make Section 3 stronger on the cap rate compression thesis."
                    ),
                    key="ir_refinement",
                )
                if st.button(
                    "🔄  Regenerate with Changes",
                    type="primary",
                    disabled=not refinement_feedback.strip(),
                    key="ir_refine_btn",
                ):
                    # Write refinement to temp file
                    rf = tempfile.NamedTemporaryFile(
                        mode="w", suffix=".txt", delete=False,
                        dir=ROOT / "configs", encoding="utf-8",
                    )
                    rf.write(refinement_feedback.strip())
                    rf.close()

                    refine_flags = ["--refinement-file", rf.name]
                    # Re-use previously selected reports (no re-extract)
                    if selected_reports:
                        refine_flags += ["--reports"] + selected_reports
                    # Pass analyst notes too if they exist
                    if analyst_notes.strip():
                        nf = tempfile.NamedTemporaryFile(
                            mode="w", suffix=".txt", delete=False,
                            dir=ROOT / "configs", encoding="utf-8",
                        )
                        nf.write(analyst_notes.strip())
                        nf.close()
                        refine_flags += ["--notes-file", nf.name]
                    else:
                        nf = None

                    _t0 = time.time()
                    ok  = _run_script("generate_investment_rationale.py",
                                      config_path, refine_flags)
                    _elapsed = time.time() - _t0
                    try: os.unlink(rf.name)
                    except: pass
                    if nf:
                        try: os.unlink(nf.name)
                        except: pass
                    if ok:
                        if _SHOW_RATIONALE_TIMING:
                            _meta_out = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
                            try:
                                import json as _jt, datetime as _dt
                                (_meta_out / "Investment_Rationale_meta.json").write_text(
                                    _jt.dumps({
                                        "generated_at": _dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
                                        "elapsed_s":    round(_elapsed, 1),
                                    }), encoding="utf-8")
                            except Exception:
                                pass
                        st.success("✅  Refined rationale generated!")
                        st.rerun()

            # ── Source audit preview ──────────────────────────────────────────
            with st.expander("🔍  Source Audit Preview", expanded=False):
                st.caption(
                    "Every cited data point, its source, and backend verification status. "
                    "🔴 Red rows need manual PDF verification. "
                    "Download the Excel for the full review workflow."
                )
                try:
                    import openpyxl
                    wb_audit = openpyxl.load_workbook(audit_file, data_only=True)
                    ws_audit = wb_audit.active
                    # Read from row 3 (headers) onwards
                    rows_audit = [
                        [str(c.value or "") for c in row]
                        for row in ws_audit.iter_rows(min_row=3)
                        if any(c.value for c in row)
                    ]
                    if len(rows_audit) > 1:
                        df_audit = pd.DataFrame(rows_audit[1:], columns=rows_audit[0])
                        st.dataframe(df_audit, hide_index=True,
                                     use_container_width=True,
                                     height=min(500, 40 + len(df_audit) * 40))
                except Exception:
                    st.info("Download the Excel file to review the source audit.")


# ─────────────────────────────────────────────────────────────────────────────
# ROUTE A2:  Existing Deals  →  view / edit / delete
# ─────────────────────────────────────────────────────────────────────────────

def render_existing_deals(deal_name: str | None, config_path: str | None):
    """
    ROUTE A2 — View, edit, or delete an existing deal.

    The user picks a deal from the sidebar radio list.  Two tabs are shown:

    📊 Summary
        Key metrics row, Property Details + Market & Search cards, output files
        (Excel previews + map PNGs), and a shortcut button to Comparable Analysis.

    ✏️ Edit
        Editable data table (same _fields_to_df editor used in New Deal Step 2).
        Saves changes in-place to the deal config JSON, preserving all other
        config sections (llm, mapbox, openai, etc.).

        Danger Zone at the bottom: two-step delete confirmation.
        First click shows a warning; second click removes the config JSON file.
    """
    st.title("📁  Existing Deals")

    if not deals:
        st.info("No deals yet — create one in **New Deal** first.")
        return

    if not deal_name or config_path is None:
        st.info("Select a deal from the sidebar.")
        return

    cfg  = load_config(config_path)
    subj = cfg["subject_property"]

    gfa_unit = subj.get("gfa_unit", "sf").upper()
    currency = subj.get("currency_symbol", subj.get("currency", ""))

    st.subheader(subj.get("deal_name", deal_name))
    st.caption(f"📍  {subj.get('address', '')}")
    st.divider()

    tab_summary, tab_edit = st.tabs(["📊  Summary", "✏️  Edit"])

    # ── Summary tab ──────────────────────────────────────────────────────────
    with tab_summary:
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Asset Class",    subj.get("asset_class", "—").title())
        m2.metric(f"GFA ({gfa_unit})", f"{int(subj.get('gfa_sf', 0)):,}")
        m3.metric("Quality",        subj.get("quality", "—"))
        m4.metric("Location",       subj.get("location", "—"))
        m5.metric("Sale Date",      subj.get("sale_date", "—"))

        st.write("")

        col_prop, col_mkt = st.columns(2)

        with col_prop:
            st.markdown("#### 🏢  Property Details")
            prop_data = {
                "Property Name":    subj.get("property_name", "—"),
                "Asset Type":       subj.get("asset_type", "—"),
                "Land Zoning":      subj.get("land_zoning", "—"),
                "Leasehold (yrs)":  subj.get("remaining_leasehold_yrs", 0),
                "Price (millions)": (f"{currency} {subj['price_sgd_m']:.1f}M"
                                     if subj.get("price_sgd_m") else "—"),
                "Cap Rate":         (f"{subj['ftm_noi_cap_rate']*100:.2f}%"
                                     if subj.get("ftm_noi_cap_rate") else "—"),
            }
            st.dataframe(pd.DataFrame(prop_data.items(), columns=["Field", "Value"]),
                         hide_index=True, use_container_width=True)

        with col_mkt:
            st.markdown("#### 🌍  Market & Search")
            kws = subj.get("submarket_keywords", [])
            if isinstance(kws, list):
                kws = ", ".join(kws)
            mkt_data = {
                "Country":            subj.get("country_name", "—"),
                "Currency":           f"{subj.get('currency','')} ({subj.get('currency_symbol','')})",
                "GFA Unit":           gfa_unit,
                "Search Keyword":     subj.get("asset_search_keyword", "—"),
                "Submarket Keywords": kws or "—",
                "Broader Query":      subj.get("broader_market_query", "—"),
            }
            st.dataframe(pd.DataFrame(mkt_data.items(), columns=["Field", "Value"]),
                         hide_index=True, use_container_width=True)

        # Output files
        out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent
        if out_dir.exists():
            xlsx_files = [f for f in sorted(out_dir.glob("*.xlsx"))
                          if not f.name.startswith("~")]
            png_files  = sorted(out_dir.glob("*.png"))

            if xlsx_files or png_files:
                st.write("")
                st.markdown("#### 📂  Output Files")
                for f in xlsx_files:
                    with st.expander(f"📊  {f.name}", expanded=False):
                        preview_df = _read_excel_preview(str(f))
                        if preview_df is not None and not preview_df.empty:
                            st.dataframe(preview_df, hide_index=True,
                                         use_container_width=True,
                                         height=min(400, 40 + len(preview_df) * 38))
                        with open(f, "rb") as fh:
                            st.download_button(
                                f"⬇️  Download {f.name}", fh.read(), f.name,
                                "application/vnd.openxmlformats-"
                                "officedocument.spreadsheetml.sheet",
                                key=f"ed_dl_{f.name}",
                                use_container_width=True)
                for f in png_files:
                    with st.expander(f"🗺️  {f.name}", expanded=False):
                        st.image(str(f), use_container_width=True)
                        with open(f, "rb") as fh:
                            st.download_button(
                                f"⬇️  Download {f.name}", fh.read(), f.name,
                                "image/png", key=f"ed_dl_{f.name}",
                                use_container_width=True)

        st.divider()
        if st.button("📊  Open in Analysis Output →", type="primary",
                     key="ed_goto_comps"):
            st.session_state["active_page"] = "overview"
            st.session_state["comp_deal"]   = deal_name
            st.rerun()

    # ── Edit tab ─────────────────────────────────────────────────────────────
    with tab_edit:
        edit_subj = dict(subj)
        if "country_code" not in edit_subj:
            edit_subj["country_code"] = cfg.get("country_code", "")

        edited = st.data_editor(
            _fields_to_df(edit_subj),
            column_config={
                "Section": st.column_config.TextColumn(
                    "Section", width="small", disabled=True),
                "Field":   st.column_config.TextColumn(
                    "Field",   width="medium", disabled=True),
                "_key":    None,
                "Value":   st.column_config.TextColumn(
                    "Value",   width="large"),
            },
            hide_index=True,
            use_container_width=True,
            height=700,
            key="ed_deal_editor",
        )

        sv_col, _ = st.columns([1, 5])
        with sv_col:
            if st.button("💾  Save Changes", key="ed_deal_save",
                         type="primary", use_container_width=True):
                updated = _df_to_fields(edited)
                cfg["subject_property"].update(updated)
                cfg["country_code"] = updated.get(
                    "country_code", cfg.get("country_code", ""))
                try:
                    Path(config_path).write_text(
                        json.dumps(cfg, indent=2, ensure_ascii=False),
                        encoding="utf-8")
                    st.success(f"✅  Saved changes to **{deal_name}**")
                except Exception as e:
                    st.error(f"Save failed: {e}")

        st.divider()
        st.markdown("#### ⚠️  Danger Zone")
        _confirm_key = f"_confirm_delete_{deal_name}"
        if not st.session_state.get(_confirm_key):
            if st.button("🗑️  Delete this deal", key=f"ed_del_btn_{deal_name}"):
                st.session_state[_confirm_key] = True
                st.rerun()
        else:
            st.warning(
                f"This will permanently delete **{deal_name}** and its config file. "
                "This cannot be undone.")
            confirm_col, cancel_col, _ = st.columns([1, 1, 4])
            with confirm_col:
                if st.button("✅  Yes, delete", key=f"ed_del_confirm_{deal_name}",
                             type="primary"):
                    try:
                        Path(config_path).unlink()
                        st.session_state.pop(_confirm_key, None)
                        st.session_state.pop("deal_nav", None)
                        st.success(f"Deal **{deal_name}** deleted.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Delete failed: {e}")
            with cancel_col:
                if st.button("❌  Cancel", key=f"ed_del_cancel_{deal_name}"):
                    st.session_state.pop(_confirm_key, None)
                    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# ROUTE:  Deal Overview  →  aggregated read-only preview + orchestrator
# ─────────────────────────────────────────────────────────────────────────────

def render_analysis_output(view: str = "overview"):
    """
    ROUTE — Analysis Output workspace.

    One deal-centric page.  A single deal selector at the top; the ``view`` (chosen
    from the sidebar's Analysis Output group) decides which of three panels shows:
      • overview   — read-only aggregated preview of every output (Sales / Rent /
                     Land comp tables + maps, and the Investment Rationale).
                     Sections with no output are skipped.
      • comps      — the editable Comparable Analysis detail page.
      • rationale  — the Investment Rationale generate / refine detail page.

    Edits made in the detail views are saved to output/<Deal>/, and the Overview
    view re-reads those files on every rerun — so any change reflects on the
    Overview preview automatically.
    """
    st.title("📊  Analysis Output")

    if not deals:
        st.info("No deals yet — create one in **New Deal** first.")
        return

    # Single deal selector drives every view. Seeded from / written to the shared
    # ``comp_deal`` key so it stays in sync across the workspace.
    _keys = list(deals.keys())
    _default_idx = 0
    if st.session_state.get("comp_deal") in deals:
        _default_idx = _keys.index(st.session_state["comp_deal"])
    deal = st.selectbox("Select Deal", _keys, index=_default_idx,
                        key="ao_deal_select")
    st.session_state["comp_deal"] = deal
    config_path = deals[deal]
    cfg  = load_config(config_path)
    subj = cfg["subject_property"]

    st.caption(f"📍  {subj.get('address') or '—'}")
    st.divider()

    if view == "comps":
        render_comparable_analysis(deal, config_path)
    elif view == "rationale":
        render_investment_rationale(deal, config_path)
    else:                                 # overview
        _render_overview_preview(deal, config_path, cfg)


_OV_COMP_ROUTES = {   # label → (excel_key, pdf_key, image_key, scan_script, output_prefix)
    "📄  Rent":        ("rent_input_file", "rent_input_pdf_file", "rent_input_image_file",
                        "scan_input_rent_comps.py", "Rent_Comps"),
    "🏢  Asset Sales": ("input_file", "input_pdf_file", "input_image_file",
                        "scan_input_sales_comps.py", "Transaction_Comparables"),
    "🌍  Land":        ("land_input_file", "land_input_pdf_file", "land_input_image_file",
                        "scan_input_land_comps.py", "Land_Sale_Comps"),
}

# output prefix → its three config input keys (derived from the routes above)
_PREFIX_INPUT_KEYS = {r[4]: (r[0], r[1], r[2]) for r in _OV_COMP_ROUTES.values()}

# upload output prefix → (online-search script, online output prefix)
_ONLINE_ROUTES = {
    "Rent_Comps":              ("search_online_rent_comps.py",  "Online_Rent_Comps"),
    "Transaction_Comparables": ("search_online_sales_comps.py", "Online_Comparables"),
    "Land_Sale_Comps":         ("search_online_land_comps.py",  "Online_Land_Comps"),
}
# upload prefix → scan script (for the Generate-all step builder)
_SCAN_BY_PREFIX = {r[4]: r[3] for r in _OV_COMP_ROUTES.values()}


def _latest_comp_excel(out_dir, upload_prefix):
    """Newest comp Excel for a type across upload + online-search prefixes.
    Returns (Path, prefix_used) or (None, None) — so the Overview shows whichever
    source (uploaded scan or online search) was produced most recently."""
    prefixes = [upload_prefix]
    if upload_prefix in _ONLINE_ROUTES:
        prefixes.append(_ONLINE_ROUTES[upload_prefix][1])
    files = []
    for _pfx in prefixes:
        files += [f for f in out_dir.glob(f"{_pfx}*.xlsx") if not f.name.startswith("~")]
    if not files:
        return None, None
    latest = max(files, key=lambda f: f.stat().st_mtime)
    used = next((p for p in prefixes if latest.name.startswith(p)), upload_prefix)
    return latest, used


_PREFIX_TO_TYPE = {
    "Transaction_Comparables": "sales",
    "Rent_Comps":              "rent",
    "Land_Sale_Comps":         "land",
}

# comp type ("sales"/"rent"/"land") → (label, excel_key, pdf_key, image_key, script, prefix)
_TYPE_TO_COMP = {_PREFIX_TO_TYPE[_r[4]]: (_lbl, *_r)
                 for _lbl, _r in _OV_COMP_ROUTES.items()}


def _classify_uploads(files):
    """Auto-detect the comp type (sales / rent / land) of each staged upload using
    backend/comp_classifier.py. Keyword-only (instant, offline) — the analyst can
    override in the UI. Cached per (name,size) so Streamlit reruns don't reclassify.
    Returns [{name, type, label, confidence, reason, scores}]."""
    import comp_classifier as clf
    cache = st.session_state.setdefault("_ov_cls_cache", {})
    out = []
    for uf in files:
        data = uf.getvalue()
        key = f"{uf.name}:{len(data)}"
        if key not in cache:
            suf = Path(uf.name).suffix or ".bin"
            tmp = tempfile.NamedTemporaryFile(suffix=suf, delete=False)
            try:
                tmp.write(data)
                tmp.close()
                res = clf.classify_file(tmp.name, allow_llm=False)
            except Exception as _e:
                res = {"type": "unknown", "label": "Unknown", "confidence": "none",
                       "reason": f"could not read ({_e})", "scores": {}}
            finally:
                try:
                    os.unlink(tmp.name)
                except Exception:
                    pass
            res["name"] = uf.name          # show the real name, not the temp path
            cache[key] = res
        out.append(cache[key])
    return out


def _comp_summary_for_rationale(out_dir):
    """Compact comparable-evidence text for the rationale writer, built from the
    SAME PGIM grids shown on the Overview (exact numbers, incl. the computed Average
    row). Passed via --comps-file so the write-up can benchmark the subject's pricing
    against the comps. Returns '' when no comp output exists."""
    blocks = []
    for prefix, _heading, title, _sb, _cb, avg_kw in _COMP_TYPES:
        latest, _used = _latest_comp_excel(out_dir, prefix)
        if not latest:
            continue
        grid = _read_pgim_grid(str(latest), avg_kw)
        if not grid or not grid.get("comps"):
            continue
        header = grid["header"]
        lines = [f"### {title} — {len(grid['comps'])} comparable(s)",
                 "Columns: " + " | ".join(h for h in header if h)]
        for srow in (grid.get("subject") or []):
            if any(srow):
                lines.append("SUBJECT: " + " | ".join(srow))
        for i, row in enumerate(grid["comps"], 1):
            lines.append(f"{i}. " + " | ".join(row))
        if grid.get("average"):
            lines.append("AVERAGE: " + " | ".join(grid["average"]))
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def _run_comp_agent(comp_type, prefix, config_path, cfg, has_files, online_enabled,
                    max_iters=3):
    """Bounded acquire→verify→evaluate→reflect→fallback loop for one comp type
    (see docs/comp_acquisition_agent.md). Runs the scan / online-search scripts
    as tools via _run_script, then uses backend/comp_acquisition_agent.py to grade
    the result and pick the next source. Emits live st.write lines (call inside an
    st.status). Returns (summary, trace)."""
    import comp_acquisition_agent as agent
    out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent

    _online = _ONLINE_ROUTES.get(prefix)
    _srcmap = {                                    # source → (script, output prefix)
        "files":  (_SCAN_BY_PREFIX.get(prefix), prefix),
        "online": (_online[0] if _online else None, _online[1] if _online else prefix),
    }
    plan = []
    if has_files:
        plan.append("files")
    if online_enabled or not has_files:            # online: chosen, or the only source
        plan.append("online")
    plan = [s for s in plan if _srcmap[s][0]][:max_iters]
    if not plan:
        return None, []

    _full = load_config(config_path)
    _full["_root"] = str(ROOT)
    llm_cfg    = _full.get("llm", {})
    openai_key = os.environ.get("OPENAI_API_KEY", "")

    tried, trace, summary = [], [], None
    source = plan[0]
    while source and len(tried) < max_iters:
        script, src_prefix = _srcmap[source]
        st.write(f"🔄 **{comp_type.title()} — {source}** · `{script}`")
        try:
            _run_script(script, config_path, ["--map"], stream_live=False)
        except Exception as _e:
            st.write(f"   ⚠ run error: {_e}")
        run_log = st.session_state.get(f"_runlog_{script}", "")
        records = agent.read_records(out_dir, src_prefix)
        src_txt = agent.source_text(_full, comp_type, source, out_dir)
        checks  = agent.verify_comps(records, src_txt, comp_type)
        ev      = agent.evaluate(records, checks, run_log, comp_type)
        tried.append(source)
        trace.append({"source": source, **ev})
        summary = {"source_used": source, "prefix": src_prefix, **ev}
        st.write(f"   → {ev['n_valid']} valid · {int(ev['pct_grounded'] * 100)}% grounded "
                 f"· confidence {ev['confidence']}"
                 + (f" · ⚑ {', '.join(ev['flags'])}" if ev['flags'] else ""))
        if ev["ok"]:
            break
        remaining = [s for s in plan if s not in tried and _srcmap[s][0]]
        refl = agent.reflect(ev, tried, remaining, llm_cfg, openai_key)
        trace.append({"reflect": refl})
        st.write(f"   🤔 {refl['diagnosis']}")
        if refl["next_action"] == "stop" or refl["next_action"] not in remaining:
            break
        source = refl["next_action"]
    return summary, trace


def _render_overview_preview(deal: str, config_path: str, cfg: dict):
    """The Overview: two upload boxes at the top (comparable files + market
    reports) drive the agentic pipeline, then a STATIC PGIM-standard view of the
    results — a colored table + location map for Sales / Land / Rent (each shown
    only if it exists), followed by the Investment Rationale, and a one-click
    download of everything as a single Word file. Editing happens in the detail
    views; new uploads here (and there) APPEND to the deal's existing inputs."""
    out_dir = ROOT / Path(cfg.get("output_file", "output/x/x.xlsx")).parent

    # ── Fresh browser session → reset this deal's Overview ───────────────────
    # Drop previously-saved comp inputs from the config and delete this deal's
    # generated outputs, so nothing carries over between sessions (Overview
    # uploads are treated as session-only). Runs once per deal per session.
    _init_key = f"_ov_init_{deal}"
    if not st.session_state.get(_init_key):
        st.session_state[_init_key] = True
        _c = load_config(config_path)
        _changed = False
        for _keys in _PREFIX_INPUT_KEYS.values():
            for _k in _keys:
                if _c.pop(_k, None) is not None:
                    _changed = True
        if _changed:
            Path(config_path).write_text(
                json.dumps(_c, indent=2, ensure_ascii=False), encoding="utf-8")
            cfg = _c
        if out_dir.exists():
            for _pref in ("Transaction_Comparables", "Rent_Comps", "Land_Sale_Comps",
                          "Online_Comparables", "Online_Rent_Comps", "Online_Land_Comps",
                          "Investment_Rationale", "Source_Audit"):
                for _of in out_dir.glob(f"{_pref}*"):
                    try:
                        _of.unlink()
                    except Exception:
                        pass

    # Market reports are shared across deals — clear the library once per session
    # (a session-global flag, so opening a second deal doesn't wipe reports you
    # just uploaded for the first one).
    if not st.session_state.get("_ov_reports_cleared"):
        st.session_state["_ov_reports_cleared"] = True
        _rdir0 = ROOT / "Input_files" / "market_reports"
        if _rdir0.exists():
            for _pdf in _rdir0.glob("*.pdf"):
                try:
                    _pdf.unlink()
                except Exception:
                    pass

    subj = cfg["subject_property"]

    # Inline subject info strip
    i1, i2, i3, i4 = st.columns(4)
    i1.markdown(f"**Class:** {(subj.get('asset_class') or '—').title()}")
    i2.markdown(f"**GFA:** {int(subj.get('gfa_sf', 0) or 0):,} "
                f"{subj.get('gfa_unit', 'sf').upper()}")
    i3.markdown(f"**Quality:** {subj.get('quality') or '—'}")
    i4.markdown(f"**Country:** {subj.get('country_name') or '—'}")

    # ════════════════════════════════════════════════════════════════════════
    # 1 · INPUTS — save each type's files (distinct), then Generate all at once
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("#### 1 · Provide inputs")
    box_c, box_r = st.columns(2)

    # ── Comparable files — pick a type, upload, SAVE (persist only, no run).
    #    Per-type uploader key means files never bleed across types. ──────────
    with box_c:
        with st.container(border=True):
            st.markdown("**📁 Comparable files** — drop mixed Excel / PDF / image; "
                        "the classifier auto-sorts them into Rent / Asset Sales / Land")
            ov_comp_files = st.file_uploader(
                "Upload comparable files (any type)",
                type=["xlsx", "pdf", "png", "jpg", "jpeg"],
                accept_multiple_files=True, key="ov_up_multi",
                label_visibility="collapsed")

            # Auto-detect each staged file's comp type; let the analyst override.
            _assign = []                     # [(uploaded_file, comp_type)]
            _labels_by_type = {t: v[0] for t, v in _TYPE_TO_COMP.items()}
            _type_labels = list(_labels_by_type.values())
            if ov_comp_files:
                _cls = {c["name"]: c for c in _classify_uploads(ov_comp_files)}
                st.caption("🪄 Detected type per file — override if wrong:")
                for _uf in ov_comp_files:
                    _c = _cls.get(_uf.name, {})
                    _dlabel = _labels_by_type.get(_c.get("type", "unknown"))
                    _idx = _type_labels.index(_dlabel) if _dlabel in _type_labels else 0
                    _badge = ("✅" if _c.get("confidence") == "high"
                              else "🟡" if _c.get("confidence") == "low" else "❓")
                    _cc1, _cc2 = st.columns([3, 2])
                    _cc1.markdown(f"{_badge} `{_uf.name}`")
                    _cc1.caption(_c.get("reason", ""))
                    _sel = _cc2.selectbox(
                        "type", _type_labels, index=_idx,
                        key=f"ov_sort_{_uf.name}", label_visibility="collapsed")
                    _seltype = next(t for t, lb in _labels_by_type.items() if lb == _sel)
                    _assign.append((_uf, _seltype))

            if st.button("💾  Save & sort", key="ov_comp_save",
                         use_container_width=True, disabled=not ov_comp_files):
                _c = load_config(config_path)
                _byt = {}
                for _uf, _t in _assign:
                    _byt.setdefault(_t, []).append(_uf)
                for _t, _grp in _byt.items():
                    _lbl, _ek, _pk, _ik, _scr, _pfx = _TYPE_TO_COMP[_t]
                    _c = _save_uploads_append(_c, _ek, _pk, _ik, _grp)
                Path(config_path).write_text(
                    json.dumps(_c, indent=2, ensure_ascii=False), encoding="utf-8")
                st.success("Saved & sorted (by comp_classifier): "
                           + ", ".join(f"{len(g)}→{_TYPE_TO_COMP[t][0].strip()}"
                                       for t, g in _byt.items()))
                st.rerun()

            # Saved files, grouped by type, each independently removable.
            for _t, (_lbl, _ek, _pk, _ik, _scr, _pfx) in _TYPE_TO_COMP.items():
                _saved_c = (_input_list(cfg.get(_ek)) + _input_list(cfg.get(_pk))
                            + _input_list(cfg.get(_ik)))
                if not _saved_c:
                    continue
                st.caption(f"✅ **{_lbl.strip()}**: "
                           + ", ".join(f"`{Path(s).name}`" for s in _saved_c))
                with st.expander(f"🗑️  Remove {_lbl.strip()} files"):
                    _rm = st.multiselect(
                        "Files to remove", [Path(s).name for s in _saved_c],
                        key=f"ov_rm_{_ek}", label_visibility="collapsed",
                        placeholder="Choose file(s) to remove…")
                    if st.button("Remove selected", key=f"ov_rm_btn_{_ek}",
                                 disabled=not _rm, use_container_width=True):
                        _c = load_config(config_path)
                        for _k in (_ek, _pk, _ik):
                            _lst = [p for p in _input_list(_c.get(_k))
                                    if Path(p).name not in _rm]
                            if _lst:
                                _c[_k] = _lst[0] if len(_lst) == 1 else _lst
                            else:
                                _c.pop(_k, None)
                        Path(config_path).write_text(
                            json.dumps(_c, indent=2, ensure_ascii=False),
                            encoding="utf-8")
                        # If no inputs remain for this type, clear its stale output.
                        if not any(_input_list(_c.get(_k)) for _k in (_ek, _pk, _ik)):
                            for _of in out_dir.glob(f"{_pfx}*"):
                                try:
                                    _of.unlink()
                                except Exception:
                                    pass
                        st.success(f"Removed {len(_rm)} file(s) from {_lbl.strip()}.")
                        st.rerun()

            # Web fallback — per type (used when a type's files are missing / weak).
            with st.expander("🌐 Web fallback per type"):
                st.caption("If a type's files don't yield enough grounded comps (or "
                           "no files were uploaded for it), the agent searches online "
                           "using the deal's keywords (needs an OpenAI key).")
                for _t, (_lbl, _ek, _pk, _ik, _scr, _pfx) in _TYPE_TO_COMP.items():
                    st.checkbox(f"Allow web fallback for {_lbl.strip()}",
                                key=f"ov_online_{_pfx}")

    # ── Market reports — upload, SAVE (persist only, no run) ─────────────────
    with box_r:
        with st.container(border=True):
            st.markdown("**📄 Market reports** — PDF (for the rationale)")
            ov_rpt_files = st.file_uploader(
                "Upload reports", type=["pdf"], accept_multiple_files=True,
                key="ov_rpt_up", label_visibility="collapsed")
            _rdir = ROOT / "Input_files" / "market_reports"
            _saved_r = sorted(_rdir.glob("*.pdf")) if _rdir.exists() else []
            if _saved_r:
                st.caption("✅ Saved reports: "
                           + ", ".join(f"`{p.name}`" for p in _saved_r))
                with st.expander("🗑️  Remove saved reports"):
                    st.caption("Removes the PDF from the shared market-reports "
                               "library (affects all deals).")
                    _rmr = st.multiselect(
                        "Reports to remove", [p.name for p in _saved_r],
                        key="ov_rm_rpt", label_visibility="collapsed",
                        placeholder="Choose report(s) to remove…")
                    if st.button("Remove selected", key="ov_rm_rpt_btn",
                                 disabled=not _rmr, use_container_width=True):
                        for _p in _saved_r:
                            if _p.name in _rmr:
                                try:
                                    _p.unlink()
                                except Exception:
                                    pass
                        st.success(f"Removed {len(_rmr)} report(s).")
                        st.rerun()
            if st.button("💾  Save reports", key="ov_rpt_save",
                         use_container_width=True, disabled=not ov_rpt_files):
                _rdir.mkdir(parents=True, exist_ok=True)
                for _uf in ov_rpt_files:
                    (_rdir / _uf.name).write_bytes(_uf.getvalue())
                st.success(f"Saved {len(ov_rpt_files)} report(s).")
                st.rerun()

    # ── Generate ALL — the ORCHESTRATOR plans which agent + tool runs each task,
    #    then drives them: comp-acquisition agents per type, then the rationale
    #    writer (enriched with the comparable evidence just produced). ──────────
    import orchestrator as orch
    _rdir = ROOT / "Input_files" / "market_reports"
    _has_reports = _rdir.exists() and any(_rdir.glob("*.pdf"))
    _comp_inputs  = {_PREFIX_TO_TYPE[_p]: any(cfg.get(k) for k in _PREFIX_INPUT_KEYS[_p])
                     for _p in _PREFIX_TO_TYPE}
    _online_flags = {_PREFIX_TO_TYPE[_p]: bool(st.session_state.get(f"ov_online_{_p}"))
                     for _p in _PREFIX_TO_TYPE}
    _plan = orch.build_plan(comp_inputs=_comp_inputs, online_flags=_online_flags,
                            has_reports=_has_reports)

    st.write("")
    if _plan:
        with st.expander("📋 Orchestration plan — which agent runs which tool",
                         expanded=True):
            st.caption("The orchestrator routes each task to an agent + a "
                       "deterministic tool, in order:")
            for _ln in orch.describe_plan(_plan):
                st.markdown(f"- {_ln}")
    else:
        st.caption("Save at least one comparable type or market reports above, "
                   "then Generate.")

    if st.button("▶  Generate all", type="primary", key="ov_generate_all",
                 use_container_width=True, disabled=not _plan):
        with st.status("Running analysis…", expanded=True) as _st:
            for _step in _plan:
                if _step["task"] == "acquire":
                    _t = _step["type"]
                    _pfx = _TYPE_TO_COMP[_t][5]
                    st.markdown(f"**{_step['title']}** · _{_step['agent']}_")
                    try:
                        _summary, _ = _run_comp_agent(
                            _t, _pfx, config_path, cfg,
                            bool(_comp_inputs.get(_t)), bool(_online_flags.get(_t)))
                        if _summary and _summary.get("confidence", 1) < 0.5:
                            st.warning(f"⚠ Low confidence ({_summary['confidence']}) "
                                       f"for {_step['title']} — please review before use.")
                    except Exception as _e:
                        st.error(f"{_step['title']} failed: {_e}")
                elif _step["task"] == "write_rationale":
                    st.markdown(f"**{_step['title']}** · _{_step['agent']}_")
                    # Feed the comps just produced into the write-up (via --comps-file).
                    _cf = []
                    try:
                        _cs = _comp_summary_for_rationale(out_dir)
                        if _cs:
                            out_dir.mkdir(parents=True, exist_ok=True)
                            _cs_path = out_dir / "_comp_summary.txt"
                            _cs_path.write_text(_cs, encoding="utf-8")
                            _cf = ["--comps-file", str(_cs_path)]
                            st.write("   📎 Injecting comparable pricing evidence "
                                     "into the write-up")
                    except Exception as _e:
                        st.write(f"   ⚠ comp-summary skipped: {_e}")
                    try:
                        _run_script("generate_investment_rationale.py", config_path,
                                    _cf, stream_live=False)
                    except Exception as _e:
                        st.error(f"Investment Rationale failed: {_e}")
            _st.update(label="✅ Analysis complete — output updated below.",
                       state="complete", expanded=True)
        # No st.rerun(): the static output section below re-reads the freshly
        # written files on this same run, so the status above stays visible.

    st.divider()

    # ════════════════════════════════════════════════════════════════════════
    # 2 · OUTPUT — static PGIM-standard view
    # ════════════════════════════════════════════════════════════════════════
    st.markdown("#### 2 · Output")
    _any_shown = False
    for _prefix, _heading, _title, _sub_banner, _comp_banner, _avg_kw in _COMP_TYPES:
        _latest, _used = _latest_comp_excel(out_dir, _prefix)
        if not _latest:
            continue                     # skip if no info
        _any_shown = True
        st.markdown(f"### {_heading}")
        _html = _pgim_table_html(str(_latest), _sub_banner, _comp_banner, _avg_kw)
        if _html:
            st.markdown(_html, unsafe_allow_html=True)
        else:
            st.caption("(No table preview available.)")
        _maps = sorted(out_dir.glob(f"{_used}*_map.png"))
        if _maps:
            st.image(str(_maps[-1]), use_container_width=True)
        st.caption("✏️ To edit the table or map, open **📋 Comparable Analysis** "
                   "in the sidebar.")
        st.divider()

    # ── Investment Rationale ─────────────────────────────────────────────────
    _rat = out_dir / "Investment_Rationale.md"
    if _rat.exists():
        _any_shown = True
        st.markdown("### Investment Rationale")
        try:
            # Overview shows the rationale prose only — the metadata header +
            # Input Summary stay in the ✍️ subtab (which renders the full doc).
            # Escape '$' so Streamlit doesn't read "S$91.8 … S$124.6" as LaTeX
            # math (which renders italic).
            st.markdown(_rationale_body(_rat.read_text(encoding="utf-8"))
                        .replace("$", "\\$"))
        except Exception as _e:
            st.error(f"Could not read rationale: {_e}")
        st.caption("✏️ To edit or regenerate, open **✍️ Investment Rationale** "
                   "in the sidebar.")
        st.divider()

    # ── Combined download (everything in one Word file) ──────────────────────
    if _any_shown:
        try:
            _docx = _build_combined_docx(deal, cfg)
        except Exception as _e:
            _docx = None
            st.caption(f"(Combined report unavailable: {_e})")
        if _docx:
            st.download_button(
                "⬇️  Download full report (Word)", _docx,
                file_name=f"{deal.replace(' ', '_')}_IC_Report.docx",
                mime="application/vnd.openxmlformats-officedocument."
                     "wordprocessingml.document",
                use_container_width=True, key="ov_dl_docx")
            st.caption("One file with all comp tables, maps and the rationale. "
                       "(PDF export needs an extra library — ask to enable it.)")
    else:
        st.info("No outputs yet — upload comparable files and/or market reports "
                "above and click **Generate**.")


# ═════════════════════════════════════════════════════════════════════════════
# ROUTER  — dispatch to the correct render function based on sidebar nav
# ─────────────────────────────────────────────────────────────────────────────
# Streamlit re-executes this entire file on every user interaction.
# The router reads the current sidebar selection from session state and calls
# the appropriate render function.  Each render function draws its own content
# into the main area and is fully self-contained.
# ═════════════════════════════════════════════════════════════════════════════

active_page = st.session_state.get("active_page", "new_deal")

if active_page == "new_deal":
    render_new_deal_form()
elif active_page == "existing":
    deal_nav = st.session_state.get("deal_nav")
    if deal_nav and deal_nav in deals:
        render_existing_deals(deal_nav, deals[deal_nav])
    else:
        render_existing_deals(None, None)
else:                                    # overview | comps | rationale
    render_analysis_output(active_page)
